from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import shutil, os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Capa ──
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MELHORIAS PARA NOTA 10')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(10, 22, 40)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Organizado por Especialista')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(8, 145, 178)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Programa Imersao Nativa | Site de Vendas\nData: 2026-03-21')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(148, 163, 184)

doc.add_page_break()

# ── Data ──
specialists = {
    'CRO (Conversao)': {
        'color': RGBColor(239, 68, 68),
        'nota_geral': '7.8',
        'modules': [
            ('M01 Hero', '8', [
                'Adicionar "no espanhol real" no subtitulo (frase-chave obrigatoria ausente)',
                'Adicionar counter de urgencia real (timer regressivo ou numero de vagas)',
                'Trocar "Vagas limitadas" por justificativa concreta',
            ]),
            ('M02 Depoimentos 1', '7', [
                'Adicionar nome + contexto em cada card de video (ex: "Maria, 34 anos, viajou sozinha")',
                'Adicionar legenda/resumo textual abaixo de cada video',
                'Adicionar badge "Assistir depoimento" sobre o thumbnail',
            ]),
            ('M03 Dores', '9', [
                'Adicionar segundo micro-prova social (alem do Wanderlei)',
            ]),
            ('M04 Objecoes', '8', [
                'Considerar mover para depois do M11 (objecoes convertem mais apos a oferta)',
                'Adicionar micro-CTA apos cada objecao respondida (link ancora para M11)',
            ]),
            ('M05 Comparativo', '8', [
                'Adicionar prova social no card "bom": "5.000 alunos ja comprovaram"',
                'Adicionar badge "METODO COMPROVADO" no card do Imersao Nativa',
            ]),
            ('M06 Intro Pilares', '6', [
                'CRITICO: Adicionar CTA secundario: "Conheca o metodo completo"',
                'Adicionar micro-prova social: "O metodo que ja transformou 5.000 brasileiros"',
                'Expandir modulo: 3 icones representando cada pilar com 1 frase cada',
            ]),
            ('M06B Pilar 1', '8', [
                'Adicionar CTA no final: "Quero comecar minha jornada"',
                'Adicionar micro-prova em cada fase (ex: "A Eliane chegou aqui em 3 meses")',
            ]),
            ('M06C Pilar 2', '8', [
                'Adicionar CTA: "Quero conhecer o app"',
                'Adicionar screenshot/mockup real do app',
            ]),
            ('M06D Pilar 3', '7', [
                'Adicionar CTA: "Quero ter suporte 24h"',
                'Adicionar depoimento de aluno sobre o suporte',
                'Adicionar badge de tempo medio de resposta: "< 2h no WhatsApp"',
            ]),
            ('M07 Depoimentos Transf.', '8', [
                'Adicionar video em pelo menos 2 dos 4 depoimentos (atualmente so texto)',
                'Adicionar foto do aluno ao lado de cada depoimento textual',
            ]),
            ('M08 Quem Sou Eu', '8', [
                'Adicionar historia pessoal da Ale (momento marcante, nao so dados)',
                'Considerar mover para antes dos pilares (constroi autoridade antes do mecanismo)',
            ]),
            ('M09 Depoimentos Finais', '7', [
                'Adicionar nome + resultado teaser em cada card de video',
                'Adicionar mais videos (3 e pouco para esta posicao)',
                'Adicionar mini-texto de contexto sob cada video',
            ]),
            ('M10 Bonus', '8', [
                'CRITICO: Adicionar CTA no final: "Quero garantir meus bonus"',
                'Adicionar timer ou frase de escassez',
            ]),
            ('M11 Preco', '9', [
                'Adicionar timer de urgencia real abaixo do badge de desconto',
                'Adicionar comparativo: "Menos que 1 aula particular por mes"',
            ]),
            ('M12 Garantia', '8', [
                'Adicionar micro-depoimento: "Quase pedi reembolso, mas na segunda semana ja estava falando"',
                'Adicionar selo visual de garantia mais proeminente',
            ]),
            ('M13 FAQ', '7', [
                'Adicionar pergunta: "Quais sao as formas de pagamento?"',
                'Adicionar pergunta: "Funciona no celular?"',
                'Adicionar micro-prova social nas respostas',
            ]),
        ]
    },
    'Copy': {
        'color': RGBColor(245, 158, 11),
        'nota_geral': '7.5',
        'modules': [
            ('M01 Hero', '7', [
                '[FEITO] Trocar CTA generico para "QUERO DESTRAVAR MEU ESPANHOL AGORA"',
                '[FEITO] Remover "simples" (palavra proibida) do subtitulo',
                'Adicionar "como um nativo fala" ou "no espanhol real" no subtitulo',
                'Trocar "que cabe na sua rotina" por: "com 15 minutos por dia, direto do celular"',
                'Substituir "e a comunidade so cresce!" por dado concreto: "em 17 paises"',
            ]),
            ('M02 Depoimentos 1', '7', [
                'Corrigir pontuacao: "vida no exterior , cada historia" usar travessao',
                'Adicionar frase de reforco: "Mais de 5.000 historias como essas"',
                'Variar o CTA (ja usado em outro modulo)',
            ]),
            ('M03 Dores', '9', [
                '[FEITO] Corrigir "Seja honesta com voce mesma" para neutro',
                'Corrigir pontuacao sistemica: todas as " , " por travessao nos 6 cards',
                'Adicionar ponto final consistente em todos os cards',
                'Trocar "espanhol real" (saturado) por "espanhol do dia a dia"',
            ]),
            ('M04 Objecoes', '8', [
                'Corrigir pontuacao: "curtas e objetivas , feitas" usar travessao',
                'Diferenciar objecoes 2 e 3 (tempo vs volume), quase identicas',
                'Trocar "condicionado" por "condicionado(a)" para genero inclusivo',
                'Trocar CTA generico por versao com beneficio',
            ]),
            ('M05 Comparativo', '7', [
                'Trocar "QECR europeu" por "padrao europeu de niveis (do A1 ao C2)"',
                'Eliminar repeticao: "como os nativos realmente falam" aparece 2x',
                'Fundir itens 3 e 4 da coluna positiva (quase identicos)',
                'Trocar CTA por versao mais forte: "Quero o metodo que funciona"',
                'Adicionar voz da Ale (bloco soa institucional)',
            ]),
            ('M06 Intro Pilares', '7', [
                'Reescrever subtitulo com frase mais concreta',
                'Adicionar voz da Ale: "Eu montei esse programa em cima de 3 pilares porque..."',
            ]),
            ('M06B Pilar 1', '7', [
                'Trocar "fluencia" por "confianca" (pode soar como promessa proibida)',
                'Substituir habilidades genericas por exemplos tangiveis',
                'Reduzir repeticoes: "de verdade" 2x, "naturalidade" 2x',
                'Adicionar tom da Ale nos intros de cada fase (atualmente tecnico)',
            ]),
            ('M06C Pilar 2', '7', [
                'Trocar "Sistema de repeticao espacada" por linguagem leiga',
                'Reduzir saturacao: "nativo" 3x, "real" 3x, "brasileiro" 3x',
                'Trocar "Certificacao reconhecida mundialmente" por dado especifico',
                'Adicionar tom pessoal da Ale',
            ]),
            ('M06D Pilar 3', '7', [
                'Trocar "sozinho" por "sozinho(a)" em genero inclusivo',
                'Trocar "perdido" por "perdido(a)"',
                'Trocar "travado" por "travado(a)"',
                'Eliminar cliches: "A combinacao perfeita"',
                'Reduzir repeticoes: "brasileiros" 2x, "IA 24/7" 2x',
            ]),
            ('M07 Depoimentos Transf.', '9', [
                'Trocar subtitulo generico por: "Novas carreiras. Novos paises. Nova confianca."',
            ]),
            ('M08 Quem Sou Eu', '9', [
                'Trocar "como sua amiga nativa" por "como quem te entende de verdade"',
                'Trocar "voce vai ter a mim" por "eu estarei ao seu lado"',
                'Reduzir dados repetidos com M6C e M7',
            ]),
            ('M09 Depoimentos Finais', '6', [
                'Trocar "a escolha certa" (cliche) por frase original',
                'Trocar "Resultados reais de alunos reais" (redundante) por "Sem roteiro. So a verdade."',
                'Corrigir pontuacao: "destravaram , e voce?" usar travessao',
                'Trocar labels "Depoimento" pelos nomes dos alunos',
            ]),
            ('M10 Bonus', '7', [
                'Trocar "fazem toda a diferenca" por "aceleram seus primeiros 30 dias"',
                'Adicionar exemplos concretos: "De Shakira a Bad Bunny, aprenda cantando"',
                'Trocar "de forma leve" por "sem parecer aula"',
            ]),
            ('M11 Preco', '8', [
                'Corrigir badge: virgula por travessao',
                'Encurtar titulo principal (longo demais)',
                'Reescrever item 4 (construcao confusa sobre audios)',
            ]),
            ('M12 Garantia', '8', [
                'Corrigir pontuacao: virgula por travessao',
                'Especificar abertura: "Ja pagou por curso que nem terminou?"',
                'Concretizar fecho: "antes dos 7 dias voce ja vai estar falando..."',
                'Padronizar imperativo: "Testa" para "Teste"',
            ]),
            ('M13 FAQ', '7', [
                'Corrigir pontuacao: todas as " , " por travessao nas respostas',
                'Trocar respostas vagas por concretas',
                'Resumir resposta sobre garantia (repete o M12)',
                'Trocar "para sempre" por "acesso permanente"',
                'Encurtar respostas (FAQ = direto ao ponto)',
            ]),
        ]
    },
    'Design': {
        'color': RGBColor(99, 102, 241),
        'nota_geral': '7.3',
        'modules': [
            ('M01 Hero', '8', [
                'Adicionar animacao de hover com scale sutil no CTA',
                'Adicionar micro-animacao de entrada no trust bar (fade-in sequencial)',
            ]),
            ('M02 Depoimentos 1', '7', [
                'Padronizar cor do CTA de #facc15 para #fbbf24',
                'Trocar breakpoint de 640px para 768px',
                'Adicionar blob-center (ausente)',
            ]),
            ('M03 Dores', '7', [
                'Trocar eyebrow rosa/ouro para sky (#38bdf8), padronizar com 6C',
                'Padronizar breakpoint de 580px para 768px',
            ]),
            ('M04 Objecoes', '8', [
                'Padronizar breakpoint de 600px para 768px',
                'Adicionar CTA button visual (ausente)',
            ]),
            ('M05 Comparativo', '7', [
                'Padronizar cor do CTA para #fbbf24',
                'Padronizar breakpoint de 620px para 768px',
                'Padronizar eyebrow para tecnica pseudo-element',
            ]),
            ('M06 Intro Pilares', '6', [
                'CRITICO: Adicionar dot pulsante no eyebrow (ausente)',
                'Adicionar icones visuais dos 3 pilares',
                'Adicionar blob-center (apenas 1 blob sutil)',
            ]),
            ('M06B Pilar 1', '8', [
                'Garantir scroll-snap no mobile nas tabs',
                'Adicionar indicador visual de scroll horizontal',
            ]),
            ('M06C Pilar 2', '8', [
                'Adicionar CTA button visual (ausente)',
                'Adicionar animacao de contador nos stats',
            ]),
            ('M06D Pilar 3', '8', [
                'Adicionar CTA button visual',
                'Padronizar eyebrow laranja para sky',
            ]),
            ('M07 Depoimentos Transf.', '7', [
                'Trocar peso do titulo de 800 para 900',
                'Adicionar blob-center',
                'Padronizar breakpoint para 768px',
            ]),
            ('M08 Quem Sou Eu', '5', [
                'CRITICO: Trocar BG solido #0d2140 por gradient linear-gradient(165deg,...)',
                'CRITICO: Adicionar dot pulsante no eyebrow (atualmente ouro SEM animacao)',
                'Padronizar eyebrow ouro para sky (#38bdf8)',
                'Padronizar breakpoint de 760px para 768px',
            ]),
            ('M09 Depoimentos Finais', '7', [
                'Trocar peso do titulo de 800 para 900',
                'Adicionar blob-center',
                'Adicionar indicacao visual de play nos videos',
            ]),
            ('M10 Bonus', '7', [
                'CRITICO: Trocar BG solido #0d2140 por gradient',
                'Adicionar CTA button visual (ausente)',
                'Padronizar breakpoint de 820px para 768px',
            ]),
            ('M11 Preco', '7', [
                'Trocar peso do titulo de 800 para 900',
                'Adicionar segundo blob (apenas 1 central)',
                'Adicionar breakpoint 768px (apenas 480px existe)',
            ]),
            ('M12 Garantia', '8', [
                'Padronizar eyebrow emerald para sky',
                'Padronizar breakpoint de 680px para 768px',
            ]),
            ('M13 FAQ', '8', [
                'Adicionar micro-animacao de abertura nos cards',
                'Adicionar icone de categoria por pergunta',
            ]),
        ]
    },
    'SEO': {
        'color': RGBColor(16, 185, 129),
        'nota_geral': '8.5',
        'modules': [
            ('og:image e twitter:image', '6', [
                'CRITICO: Adicionar meta og:image com URL da imagem de compartilhamento',
                'Adicionar og:image:width e og:image:height',
                'Adicionar meta twitter:image',
            ]),
            ('Title tag', '9', [
                'Considerar incluir "Curso de" para SEO: "Curso de Espanhol Online para Brasileiros"',
            ]),
            ('Meta description', '9', [
                'Ja esta otima. Considerar testar variacao com CTA mais direto',
            ]),
            ('Hierarquia H1/H2/H3', '8', [
                'Trocar h3.pilar-title para h2 (semanticamente sao titulos de secao)',
                'Subtitulos dos pilares viram h3',
            ]),
            ('Schema Person', '9', [
                'Adicionar sameAs com links Instagram e YouTube da @espanholcomvoce',
            ]),
            ('Schema Course', '10', [
                'Nenhuma melhoria necessaria. Exemplar.',
            ]),
            ('Schema FAQPage', '10', [
                'Nenhuma melhoria necessaria. Exemplar.',
            ]),
            ('Lazy loading', '9', [
                'Verificar se imagens base64 inline poderiam ser externalizadas para cache',
            ]),
            ('Fonts', '7', [
                'Remover import de Playfair Display e Lato (nao usadas, ~50KB extra)',
            ]),
            ('Mobile responsivo', '8', [
                'Padronizar breakpoints individuais dos modulos para 768px e 480px',
            ]),
        ]
    },
}

# ── Render each specialist ──
for name, data in specialists.items():
    doc.add_heading(f'{name} — Nota Geral: {data["nota_geral"]}/10', level=1)

    # Summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, t in enumerate(['Modulo', 'Nota', 'Qtd Melhorias']):
        hdr[i].text = t
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)

    for mod_name, nota, items in data['modules']:
        row = table.add_row().cells
        row[0].text = mod_name
        row[1].text = f'{nota}/10'
        row[2].text = str(len(items))
        for cell in row:
            for p in cell.paragraphs:
                if p.runs:
                    p.runs[0].font.size = Pt(10)

    doc.add_paragraph('')

    # Details per module
    for mod_name, nota, items in data['modules']:
        doc.add_heading(f'{mod_name} ({nota} -> 10)', level=2)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            if item.startswith('[FEITO]'):
                run = p.add_run(item)
                run.font.bold = True
                run.font.color.rgb = RGBColor(16, 185, 129)
            elif item.startswith('CRITICO'):
                run = p.add_run(item)
                run.font.bold = True
                run.font.color.rgb = RGBColor(239, 68, 68)
            else:
                p.add_run(item)

    doc.add_page_break()

# ── Save ──
path_local = 'D:/EspanholComVoce/site-vendas/MELHORIAS_PARA_10.docx'
doc.save(path_local)
print(f'Saved: {path_local}')

path_drive = 'G:/Meu Drive/EspanholComVoce/MELHORIAS_PARA_10.docx'
os.makedirs(os.path.dirname(path_drive), exist_ok=True)
shutil.copy2(path_local, path_drive)
print(f'Copied: {path_drive}')
