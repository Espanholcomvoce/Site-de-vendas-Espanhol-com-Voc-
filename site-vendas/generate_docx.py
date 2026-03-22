from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import shutil, os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(20)
style_h1.font.color.rgb = RGBColor(10, 22, 40)

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(15)
style_h2.font.color.rgb = RGBColor(13, 33, 64)

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(12)
style_h3.font.color.rgb = RGBColor(8, 145, 178)

# ── Title page ──
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MELHORIAS PARA NOTA 10')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(10, 22, 40)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Programa Imersão Nativa — Site de Vendas')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(100, 116, 139)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Baseado na Auditoria Final de 4 Especialistas: CRO, Copy, Design, SEO\nData: 2026-03-21')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(148, 163, 184)

doc.add_page_break()

# ── Summary table ──
doc.add_heading('Resumo por Módulo', level=1)

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, t in enumerate(['Módulo', 'CRO', 'Copy', 'Design', 'Média']):
    hdr[i].text = t
    for p in hdr[i].paragraphs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

modules_data = [
    ('01 Hero', '8', '7', '8', '7.7'),
    ('02 Depoimentos 1', '7', '7', '7', '7.0'),
    ('03 Dores', '9', '9', '7', '8.3'),
    ('04 Objeções', '8', '8', '8', '8.0'),
    ('05 Comparativo', '8', '7', '7', '7.3'),
    ('06 Intro Pilares', '6', '7', '6', '6.3'),
    ('06B Pilar 1', '8', '7', '8', '7.7'),
    ('06C Pilar 2', '8', '7', '8', '7.7'),
    ('06D Pilar 3', '7', '7', '8', '7.3'),
    ('07 Depoimentos Transf.', '8', '9', '7', '8.0'),
    ('08 Quem Sou Eu', '8', '9', '5', '7.3'),
    ('09 Depoimentos Finais', '7', '6', '7', '6.7'),
    ('10 Bônus', '8', '7', '7', '7.3'),
    ('11 Preço', '9', '8', '7', '8.0'),
    ('12 Garantia', '8', '8', '8', '8.0'),
    ('13 FAQ', '7', '7', '8', '7.3'),
]

for m in modules_data:
    row = table.add_row().cells
    for i, v in enumerate(m):
        row[i].text = v
        for p in row[i].paragraphs:
            if p.runs:
                p.runs[0].font.size = Pt(10)

doc.add_paragraph('')

# ── Module details ──
modules = [
    {
        'title': 'MÓDULO 01 — HERO',
        'media': '7.7',
        'sections': [
            ('CRO', '8 → 10', [
                'Adicionar "no espanhol real" no subtítulo (frase-chave obrigatória ausente no Hero)',
                'Adicionar counter de urgência real (ex: "Últimas 47 vagas com desconto" ou timer regressivo)',
                'Trocar "Vagas limitadas" por justificativa concreta ("Turma de março com 100 vagas — 53 já preenchidas")',
            ]),
            ('Copy', '7 → 10', [
                '[FEITO] Trocar CTA genérico para "QUERO DESTRAVAR MEU ESPANHOL AGORA"',
                '[FEITO] Remover "simples" (palavra proibida) do subtítulo',
                'Adicionar "como um nativo fala" ou "no espanhol real" no subtítulo',
                'Trocar "que cabe na sua rotina" por: "com 15 minutos por dia, direto do celular"',
                'Substituir "e a comunidade só cresce!" por dado concreto: "em 17 países" ou "nota 4.9 na Hotmart"',
            ]),
            ('Design', '8 → 10', [
                'Adicionar animação de hover com scale sutil no CTA (pulse no load)',
                'Adicionar micro-animação de entrada no trust bar (fade-in sequencial)',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 02 — DEPOIMENTOS PARTE 1',
        'media': '7.0',
        'sections': [
            ('CRO', '7 → 10', [
                'Adicionar nome + contexto em cada card de vídeo (ex: "Maria, 34 anos — viajou sozinha pela Colômbia")',
                'Adicionar legenda/resumo textual abaixo de cada vídeo',
                'Adicionar badge "Assistir depoimento" sobre o thumbnail',
            ]),
            ('Copy', '7 → 10', [
                'Corrigir pontuação: "vida no exterior , cada história" → usar travessão',
                'Adicionar frase de reforço: "Mais de 5.000 histórias como essas"',
                'Variar CTA (já usado em outro módulo)',
            ]),
            ('Design', '7 → 10', [
                'Padronizar cor do CTA de #facc15 para #fbbf24',
                'Trocar breakpoint de 640px para 768px',
                'Adicionar blob-center',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 03 — DORES / IDENTIFICAÇÃO',
        'media': '8.3',
        'sections': [
            ('CRO', '9 → 10', [
                'Adicionar segundo micro-prova social (além do Wanderlei)',
            ]),
            ('Copy', '9 → 10', [
                '[FEITO] Corrigir "Seja honesta com você mesma" para neutro',
                'Corrigir pontuação sistêmica: todas as " , " → " — " nos 6 cards',
                'Adicionar ponto final consistente em todos os cards',
                'Trocar "espanhol real" por "espanhol do dia a dia" (saturação)',
            ]),
            ('Design', '7 → 10', [
                'Trocar eyebrow rosa/ouro para sky (#38bdf8)',
                'Padronizar breakpoint de 580px para 768px',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 04 — OBJEÇÕES',
        'media': '8.0',
        'sections': [
            ('CRO', '8 → 10', [
                'Considerar mover para depois do M11 (objeções convertem mais após a oferta)',
                'Adicionar micro-CTA após cada objeção respondida',
            ]),
            ('Copy', '8 → 10', [
                'Corrigir pontuação: "curtas e objetivas , feitas" → usar travessão',
                'Diferenciar objeções 2 e 3 (tempo vs volume) — quase idênticas',
                'Trocar "condicionado" por "condicionado(a)"',
                'Trocar CTA genérico por versão com benefício',
            ]),
            ('Design', '8 → 10', [
                'Padronizar breakpoint de 600px para 768px',
                'Adicionar CTA button visual',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 05 — COMPARATIVO',
        'media': '7.3',
        'sections': [
            ('CRO', '8 → 10', [
                'Adicionar prova social no card "bom": "5.000 alunos já comprovaram"',
                'Adicionar badge "MÉTODO COMPROVADO"',
            ]),
            ('Copy', '7 → 10', [
                'Trocar "QECR europeu" por "padrão europeu de níveis (do A1 ao C2)"',
                'Eliminar repetição: "como os nativos realmente falam" 2x',
                'Fundir itens 3 e 4 da coluna positiva',
                'Trocar CTA por versão mais forte',
                'Adicionar voz da Ale (bloco soa institucional)',
            ]),
            ('Design', '7 → 10', [
                'Padronizar cor do CTA para #fbbf24',
                'Padronizar breakpoint de 620px para 768px',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 06 — INTRO PILARES',
        'media': '6.3',
        'sections': [
            ('CRO', '6 → 10', [
                'Adicionar CTA secundário: "Conheça o método completo ↓"',
                'Adicionar micro-prova social: "O método que já transformou 5.000 brasileiros"',
                'Expandir módulo: 3 ícones representando cada pilar',
            ]),
            ('Copy', '7 → 10', [
                'Reescrever subtítulo com frase mais concreta',
                'Adicionar voz da Ale: "Eu montei esse programa em cima de 3 pilares porque..."',
            ]),
            ('Design', '6 → 10', [
                'CRÍTICO: Adicionar dot pulsante no eyebrow (ausente)',
                'Adicionar ícones visuais dos 3 pilares',
                'Adicionar blob-center',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 06B — PILAR 1 JORNADA',
        'media': '7.7',
        'sections': [
            ('CRO', '8 → 10', [
                'Adicionar CTA no final: "Quero começar minha jornada →"',
                'Adicionar micro-prova em cada fase',
            ]),
            ('Copy', '7 → 10', [
                'Trocar "fluência" por "confiança" (promessa proibida)',
                'Substituir habilidades genéricas por exemplos tangíveis',
                'Reduzir repetições: "de verdade" 2x, "naturalidade" 2x',
                'Adicionar tom da Ale nos intros de cada fase',
            ]),
            ('Design', '8 → 10', [
                'Garantir scroll-snap no mobile nas tabs',
                'Adicionar indicador visual de scroll horizontal',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 06C — PILAR 2 APP',
        'media': '7.7',
        'sections': [
            ('CRO', '8 → 10', [
                'Adicionar CTA: "Quero conhecer o app →"',
                'Adicionar screenshot/mockup real do app',
            ]),
            ('Copy', '7 → 10', [
                'Trocar "Sistema de repetição espaçada" por linguagem leiga',
                'Reduzir saturação: "nativo" 3x, "real" 3x, "brasileiro" 3x',
                'Trocar "Certificação reconhecida mundialmente" por dado específico',
                'Adicionar tom pessoal da Ale',
            ]),
            ('Design', '8 → 10', [
                'Adicionar CTA button visual',
                'Adicionar animação de contador nos stats',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 06D — PILAR 3 SUPORTE',
        'media': '7.3',
        'sections': [
            ('CRO', '7 → 10', [
                'Adicionar CTA: "Quero ter suporte 24h →"',
                'Adicionar depoimento de aluno sobre o suporte',
                'Adicionar badge de tempo médio de resposta',
            ]),
            ('Copy', '7 → 10', [
                'Trocar "sozinho" por "sozinho(a)" — gênero inclusivo',
                'Trocar "perdido" por "perdido(a)"',
                'Trocar "travado" por "travado(a)"',
                'Eliminar clichês: "A combinação perfeita"',
                'Reduzir repetições: "brasileiros" 2x, "IA 24/7" 2x',
            ]),
            ('Design', '8 → 10', [
                'Adicionar CTA button visual',
                'Padronizar eyebrow laranja para sky',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 07 — DEPOIMENTOS TRANSFORMAÇÃO',
        'media': '8.0',
        'sections': [
            ('CRO', '8 → 10', [
                'Adicionar vídeo em pelo menos 2 dos 4 depoimentos',
                'Adicionar foto do aluno ao lado de cada depoimento',
            ]),
            ('Copy', '9 → 10', [
                'Trocar subtítulo genérico por: "Novas carreiras. Novos países. Nova confiança."',
            ]),
            ('Design', '7 → 10', [
                'Trocar peso do título de 800 para 900',
                'Adicionar blob-center',
                'Padronizar breakpoint para 768px',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 08 — QUEM SOU EU',
        'media': '7.3',
        'sections': [
            ('CRO', '8 → 10', [
                'Adicionar história pessoal da Ale (momento marcante)',
                'Considerar mover para antes dos pilares',
            ]),
            ('Copy', '9 → 10', [
                'Trocar "como sua amiga nativa" por "como quem te entende de verdade"',
                'Trocar "você vai ter a mim" por "eu estarei ao seu lado"',
                'Reduzir dados repetidos com M6C e M7',
            ]),
            ('Design', '5 → 10', [
                'CRÍTICO: Trocar BG sólido por gradient',
                'CRÍTICO: Adicionar dot pulsante no eyebrow',
                'Padronizar eyebrow ouro para sky',
                'Padronizar breakpoint de 760px para 768px',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 09 — DEPOIMENTOS FINAIS',
        'media': '6.7',
        'sections': [
            ('CRO', '7 → 10', [
                'Adicionar nome + resultado teaser em cada card',
                'Adicionar mais vídeos (3 é pouco)',
                'Adicionar mini-texto de contexto sob cada vídeo',
            ]),
            ('Copy', '6 → 10', [
                'Trocar "a escolha certa" (clichê) por frase original',
                'Trocar "Resultados reais de alunos reais" (redundante)',
                'Corrigir pontuação: "destravaram , e" → usar travessão',
                'Trocar labels "Depoimento" pelos nomes dos alunos',
            ]),
            ('Design', '7 → 10', [
                'Trocar peso do título de 800 para 900',
                'Adicionar blob-center',
                'Adicionar indicação visual de play nos vídeos',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 10 — BÔNUS',
        'media': '7.3',
        'sections': [
            ('CRO', '8 → 10', [
                'CRÍTICO: Adicionar CTA no final',
                'Adicionar timer ou frase de escassez',
            ]),
            ('Copy', '7 → 10', [
                'Trocar "fazem toda a diferença" por "aceleram seus primeiros 30 dias"',
                'Adicionar exemplos concretos nos bônus',
                'Trocar "de forma leve" por "sem parecer aula"',
            ]),
            ('Design', '7 → 10', [
                'CRÍTICO: Trocar BG sólido por gradient',
                'Adicionar CTA button visual',
                'Padronizar breakpoint para 768px',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 11 — PREÇO',
        'media': '8.0',
        'sections': [
            ('CRO', '9 → 10', [
                'Adicionar timer de urgência real',
                'Adicionar comparativo: "Menos que 1 aula particular por mês"',
            ]),
            ('Copy', '8 → 10', [
                'Corrigir badge: vírgula → travessão',
                'Encurtar título principal',
                'Reescrever item 4 (construção confusa)',
            ]),
            ('Design', '7 → 10', [
                'Trocar peso do título de 800 para 900',
                'Adicionar segundo blob',
                'Adicionar breakpoint 768px',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 12 — GARANTIA',
        'media': '8.0',
        'sections': [
            ('CRO', '8 → 10', [
                'Adicionar micro-depoimento sobre quase-reembolso',
                'Adicionar selo de garantia mais proeminente',
            ]),
            ('Copy', '8 → 10', [
                'Corrigir pontuação: vírgula → travessão',
                'Especificar abertura: "Já pagou por curso que nem terminou?"',
                'Concretizar fecho: "antes dos 7 dias você já vai estar falando..."',
                'Padronizar imperativo: "Testa" → "Teste"',
            ]),
            ('Design', '8 → 10', [
                'Padronizar eyebrow emerald para sky',
                'Padronizar breakpoint para 768px',
            ]),
        ]
    },
    {
        'title': 'MÓDULO 13 — FAQ',
        'media': '7.3',
        'sections': [
            ('CRO', '7 → 10', [
                'Adicionar pergunta: "Quais são as formas de pagamento?"',
                'Adicionar pergunta: "Funciona no celular?"',
                'Adicionar micro-prova social nas respostas',
            ]),
            ('Copy', '7 → 10', [
                'Corrigir pontuação: " , " → " — " em todas as respostas',
                'Trocar respostas vagas por concretas',
                'Resumir resposta sobre garantia',
                'Trocar "para sempre" por "acesso permanente"',
                'Encurtar respostas (FAQ = direto)',
            ]),
            ('Design', '8 → 10', [
                'Adicionar micro-animação de abertura nos cards',
                'Adicionar ícone de categoria por pergunta',
            ]),
        ]
    },
]

for mod in modules:
    doc.add_heading(f"{mod['title']} (média {mod['media']})", level=2)

    for specialist, score, items in mod['sections']:
        doc.add_heading(f"{specialist}: {score}", level=3)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            if item.startswith('[FEITO]') or item.startswith('CRÍTICO'):
                run = p.add_run(item)
                run.font.bold = True
                if item.startswith('[FEITO]'):
                    run.font.color.rgb = RGBColor(16, 185, 129)
                else:
                    run.font.color.rgb = RGBColor(239, 68, 68)
            else:
                p.add_run(item)

    doc.add_paragraph('')

# ── SEO Global ──
doc.add_page_break()
doc.add_heading('SEO — Itens Globais para Nota 10', level=1)

seo_items = [
    ('og:image', 'AUSENTE', 'Adicionar meta og:image + og:image:width + og:image:height'),
    ('twitter:image', 'AUSENTE', 'Adicionar meta twitter:image'),
    ('Schema Person sameAs', 'AUSENTE', 'Adicionar links Instagram, YouTube da @espanholcomvoce'),
    ('Fonts não usadas', 'Playfair + Lato carregadas', 'Remover se não são usadas (~50KB extra)'),
    ('H2/H3 nos pilares', 'H3 como título', 'Trocar pilar-title para H2, subtítulos para H3'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr2 = table2.rows[0].cells
for i, t in enumerate(['Item', 'Status Atual', 'Ação Necessária']):
    hdr2[i].text = t
    hdr2[i].paragraphs[0].runs[0].font.bold = True

for item, status, action in seo_items:
    row = table2.add_row().cells
    row[0].text = item
    row[1].text = status
    row[2].text = action

# ── Summary ──
doc.add_paragraph('')
doc.add_heading('Resumo de Contagem', level=1)

summary = doc.add_table(rows=1, cols=2)
summary.style = 'Light Grid Accent 1'
shdr = summary.rows[0].cells
shdr[0].text = 'Status'
shdr[1].text = 'Quantidade'
for p in shdr[0].paragraphs + shdr[1].paragraphs:
    if p.runs: p.runs[0].font.bold = True

for label, count in [('Total de melhorias', '~120'), ('Já aplicadas', '3'), ('Críticas (impacto alto)', '12'), ('Médias', '~45'), ('Polimento', '~60')]:
    row = summary.add_row().cells
    row[0].text = label
    row[1].text = count

# ── Save ──
path_local = 'D:/EspanholComVoce/site-vendas/MELHORIAS_PARA_10.docx'
doc.save(path_local)
print(f"Saved: {path_local}")

# Copy to Google Drive
path_drive = 'G:/Meu Drive/EspanholComVoce/MELHORIAS_PARA_10.docx'
os.makedirs(os.path.dirname(path_drive), exist_ok=True)
shutil.copy2(path_local, path_drive)
print(f"Copied: {path_drive}")
