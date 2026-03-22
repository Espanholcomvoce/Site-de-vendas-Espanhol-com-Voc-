from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import shutil, os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── CAPA ──
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AUDITORIA FINAL V2')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 30, 56)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Programa Imersao Nativa — Nova Paleta de Marca')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0, 144, 190)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('4 Especialistas: CRO, Copy, Design, SEO\nData: 2026-03-22')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(148, 163, 184)

doc.add_page_break()

# ── RESUMO GERAL ──
doc.add_heading('Resumo Geral', level=1)

t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
for i, txt in enumerate(['Especialista', 'Nota', 'Foco']):
    hdr[i].text = txt
    hdr[i].paragraphs[0].runs[0].font.bold = True

for name, nota, foco in [
    ('CRO (Conversao)', '7.3 / 10', 'Funil, CTAs, prova social, urgencia'),
    ('Copy', '7.6 / 10', 'Voz da Ale, headlines, linguagem, tom'),
    ('Design', '7.2 / 10', 'Paleta, blobs, botoes, responsivo'),
    ('SEO', '8.1 / 10', 'Meta tags, schema, lazy loading, OG'),
    ('MEDIA GERAL', '7.6 / 10', ''),
]:
    row = t.add_row().cells
    row[0].text = name
    row[1].text = nota
    row[2].text = foco
    if name == 'MEDIA GERAL':
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph('')

# ── NOTAS POR MODULO ──
doc.add_heading('Notas por Modulo', level=1)

t2 = doc.add_table(rows=1, cols=6)
t2.style = 'Light Grid Accent 1'
hdr2 = t2.rows[0].cells
for i, txt in enumerate(['Modulo', 'CRO', 'Copy', 'Design', 'SEO', 'Media']):
    hdr2[i].text = txt
    hdr2[i].paragraphs[0].runs[0].font.bold = True

modules = [
    ('01 Hero', '8', '8', '8', '-', '8.0'),
    ('02 Depoimentos 1', '5', '7', '9', '-', '7.0'),
    ('03 Dores', '9', '9', '6', '-', '8.0'),
    ('04 Objecoes', '8', '8', '9', '-', '8.3'),
    ('05 Comparativo', '8', '8', '7', '-', '7.7'),
    ('06 Intro Pilares', '6', '6', '9', '-', '7.0'),
    ('06B Pilar 1', '7', '7', '9', '-', '7.7'),
    ('06C Pilar 2', '7', '6', '8', '-', '7.0'),
    ('06D Pilar 3', '6', '8', '8', '-', '7.3'),
    ('07 Depoimentos Transf', '7', '8', '7', '-', '7.3'),
    ('08 Quem Sou Eu', '8', '9', '7', '-', '8.0'),
    ('09 Depoimentos Finais', '5', '7', '7', '-', '6.3'),
    ('10 Bonus', '9', '7', '7', '-', '7.7'),
    ('11 Preco', '8', '7', '9', '-', '8.0'),
    ('12 Garantia', '9', '8', '7', '-', '8.0'),
    ('13 FAQ', '7', '8', '8', '-', '7.7'),
]

for m in modules:
    row = t2.add_row().cells
    for i, v in enumerate(m):
        row[i].text = v

doc.add_paragraph('')
doc.add_paragraph('SEO: 8.1/10 (avaliacao global, nao por modulo)')
doc.add_page_break()

# ── ESPECIALISTA 1: CRO ──
doc.add_heading('Especialista CRO — 7.3/10', level=1)

cro_data = [
    ('M01 Hero', '8', 'Headline forte, VSL presente, CTA "QUERO DESTRAVAR". Falta micro-prova social numerica.'),
    ('M02 Depoimentos 1', '5', 'CRITICO: Sem CTA. 4 videos sem nomes/contexto. Visitante quente sem botao.'),
    ('M03 Dores', '9', 'Excelente copy emocional. 6 dores reais. Reframe "o problema nao e voce" poderoso.'),
    ('M04 Objecoes', '8', '6 objecoes relevantes. CTA poderia ser mais forte. Falta urgencia.'),
    ('M05 Comparativo', '8', 'Formato visual de comparacao eficaz. Badge 5.000 alunos presente.'),
    ('M06 Intro', '6', 'CTA "CONHECA" e fraco, nao orientado a compra. Secao curta.'),
    ('M06B Pilar 1', '7', 'Tabs A1-C1 excelentes. Falta prova social dentro do pilar.'),
    ('M06C Pilar 2', '7', 'CTA "CONHECER" e fraco. Deveria ser orientado a resultado.'),
    ('M06D Pilar 3', '6', 'Imagem da Ale comentada (placeholder). Enfraquece pilar humano.'),
    ('M07 Depoimentos', '7', 'CRITICO: Sem CTA. Pico emocional sem botao. Stats bons.'),
    ('M08 Quem Sou Eu', '8', 'Historia pessoal cria empatia. Badges de autoridade bem posicionados.'),
    ('M09 Dep. Finais', '5', 'CRITICO: Sem CTA. 3 videos de depoimento sem botao.'),
    ('M10 Bonus', '9', 'Excelente stack de valor. R$1.091 gratis. Urgencia presente.'),
    ('M11 Preco', '8', 'Boa ancoragem R$1.497 por R$497. FALTA countdown timer.'),
    ('M12 Garantia', '9', '"O risco e todo meu" poderoso. CTA coerente.'),
    ('M13 FAQ', '7', '10 perguntas relevantes. Falta pergunta sobre cancelamento.'),
]

t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Light Grid Accent 1'
for i, txt in enumerate(['Modulo', 'Nota', 'Justificativa']):
    t3.rows[0].cells[i].text = txt
    t3.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for mod, nota, just in cro_data:
    row = t3.add_row().cells
    row[0].text = mod
    row[1].text = nota
    row[2].text = just

doc.add_paragraph('')
doc.add_heading('Top 5 Problemas CRO', level=2)
for item in [
    'CRITICO: 3 modulos de depoimento sem CTA (M02, M07, M09) — perda de conversao 15-25%',
    'Falta urgencia/escassez real no M11 (sem timer, sem vagas)',
    'Imagem da Ale comentada no M06D — contradiz "suporte humano"',
    'Tags HTML malformadas no M12 e M13 (<\\strong>)',
    'CTAs dos pilares com texto informativo em vez de orientado a resultado',
]:
    p = doc.add_paragraph(style='List Bullet')
    if 'CRITICO' in item:
        run = p.add_run(item)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 66, 28)
    else:
        p.add_run(item)

doc.add_page_break()

# ── ESPECIALISTA 2: COPY ──
doc.add_heading('Especialista Copy — 7.6/10', level=1)

copy_data = [
    ('M01 Hero', '8', 'Direta, sem formalidade. CTA forte. Faltam frases obrigatorias.'),
    ('M02 Depoimentos 1', '7', 'Prova social indireta. Sem CTA.'),
    ('M03 Dores', '9', 'Tom perfeito de amiga nativa. Todas as frases obrigatorias presentes.'),
    ('M04 Objecoes', '8', 'Respostas naturais. CTA generico "Quero comecar agora".'),
    ('M05 Comparativo', '8', 'Clara, sem jargao. Contraste forte.'),
    ('M06 Intro', '6', 'Tom institucional. CTA orientado a feature.'),
    ('M06B Pilar 1', '7', 'Descritiva, correta. Tom mais tecnico.'),
    ('M06C Pilar 2', '6', 'Funcional mas foco em features, nao dores.'),
    ('M06D Pilar 3', '8', 'Empatica. Genero inclusivo correto.'),
    ('M07 Depoimentos', '8', 'Historias autenticas com contexto real.'),
    ('M08 Quem Sou Eu', '9', '100% voz da Ale. "Sei onde voce trava" perfeito.'),
    ('M09 Dep. Finais', '7', 'Prova social pura. Sem CTA.'),
    ('M10 Bonus', '7', 'Proxima, generosa. Espaco antes de virgula.'),
    ('M11 Preco', '7', 'Objetiva. Lista de features.'),
    ('M12 Garantia', '8', 'Pessoal, "Palavra da Ale" autentico.'),
    ('M13 FAQ', '8', 'Respostas claras. Frases obrigatorias presentes.'),
]

t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Light Grid Accent 1'
for i, txt in enumerate(['Modulo', 'Nota', 'Justificativa']):
    t4.rows[0].cells[i].text = txt
    t4.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for mod, nota, just in copy_data:
    row = t4.add_row().cells
    row[0].text = mod
    row[1].text = nota
    row[2].text = just

doc.add_paragraph('')
doc.add_heading('Top 5 Problemas Copy', level=2)
for item in [
    'Espacos antes de virgulas em 15+ ocorrencias no site todo',
    'M06 e M06C com CTAs orientados a feature, nao a beneficio',
    'Frases obrigatorias ausentes no Hero (M01) — nenhuma das 3 aparece',
    'M06/M06B/M06C com tom institucional, perdem voz de "amiga nativa"',
    'M02 e M09 (depoimentos) sem CTA — pico emocional sem botao',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_page_break()

# ── ESPECIALISTA 3: DESIGN ──
doc.add_heading('Especialista Design — 7.2/10', level=1)

design_data = [
    ('M01 Hero', '8', 'Gradient + 3 blobs OK. Botao com gradiente amarelo, deveria ser verde solido.'),
    ('M02 Depoimentos 1', '9', 'Light correto, clean.'),
    ('M03 Dores', '6', 'BG solido sem gradient. Botao amarelo com texto escuro.'),
    ('M04 Objecoes', '9', 'Light correto, blobs suaves.'),
    ('M05 Comparativo', '7', 'Botao amarelo com texto escuro.'),
    ('M06/06B Pilares', '9', 'Completo e consistente.'),
    ('M06C Pilar 2', '8', 'Correto.'),
    ('M06D Pilar 3', '8', 'Correto.'),
    ('M07 Depoimentos', '7', 'Blobs com opacidade muito baixa (0.08-0.10).'),
    ('M08 Quem Sou Eu', '7', 'BG solido. Georgia serif fora da paleta.'),
    ('M09 Dep. Finais', '7', 'Blob roxo fora da paleta de marca.'),
    ('M10 Bonus', '7', 'BG solido sem gradient.'),
    ('M11 Preco', '9', 'Light correto, tipografia Inter.'),
    ('M12 Garantia', '7', 'Botao CTA final AMARELO ao inves de verde.'),
    ('M13 FAQ', '8', 'Correto.'),
]

t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Light Grid Accent 1'
for i, txt in enumerate(['Modulo', 'Nota', 'Justificativa']):
    t5.rows[0].cells[i].text = txt
    t5.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for mod, nota, just in design_data:
    row = t5.add_row().cells
    row[0].text = mod
    row[1].text = nota
    row[2].text = just

doc.add_paragraph('')
doc.add_heading('Top 5 Problemas Design', level=2)
for item in [
    'CRITICO: Botoes CTA inconsistentes — varios com gradiente amarelo em vez de verde solido #36c551',
    'CRITICO: 3 modulos escuros com BG solido sem gradient (M03, M08, M10)',
    'Blob roxo fora da paleta no M09',
    'Blobs do M07 com opacidade muito baixa (0.08-0.10 vs recomendado 0.15-0.20)',
    'Fontes Playfair Display e Lato carregadas sem uso (~80KB extra)',
]:
    p = doc.add_paragraph(style='List Bullet')
    if 'CRITICO' in item:
        run = p.add_run(item)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 66, 28)
    else:
        p.add_run(item)

doc.add_page_break()

# ── ESPECIALISTA 4: SEO ──
doc.add_heading('Especialista SEO — 8.1/10', level=1)

seo_data = [
    ('Title Tag', '9', '57 chars, keyword primaria presente.'),
    ('Meta Description', '8', '153 chars, beneficio + prova social.'),
    ('Hierarquia H1/H2/H3', '8', '1 H1, 11 H2, 7 H3. Correto.'),
    ('Schema Course', '9', 'Completo: nome, preco, instrutor, duracao.'),
    ('Schema FAQPage', '9', 'Perguntas e respostas estruturadas.'),
    ('Schema Person', '9', 'Alejandra Fajardo. Falta sameAs.'),
    ('Open Graph', '6', 'CRITICO: og:image AUSENTE.'),
    ('Twitter Card', '5', 'CRITICO: twitter:image AUSENTE.'),
    ('Lazy Loading', '9', 'Imagens e iframes cobertos.'),
    ('Canonical URL', '10', 'Presente e correto.'),
    ('Viewport Meta', '10', 'Correto.'),
    ('Robots Meta', '10', 'index, follow.'),
    ('Mobile Responsive', '9', '35+ media queries, sticky CTA.'),
    ('Font Loading', '5', '2 fontes carregadas sem uso.'),
    ('Tracking', '7', 'Referencia a tracking.js mas GTM/fbq nao inline.'),
]

t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Light Grid Accent 1'
for i, txt in enumerate(['Item', 'Nota', 'Justificativa']):
    t6.rows[0].cells[i].text = txt
    t6.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
for item, nota, just in seo_data:
    row = t6.add_row().cells
    row[0].text = item
    row[1].text = nota
    row[2].text = just

doc.add_paragraph('')
doc.add_heading('Top 5 Problemas SEO', level=2)
for item in [
    'CRITICO: og:image e twitter:image AUSENTES — compartilhamentos sem thumbnail',
    'Fontes Playfair Display e Lato carregadas sem uso (~80KB extra)',
    'GTM e Meta Pixel nao encontrados inline no HTML',
    'Falta twitter:site (@espanholcomvoce)',
    'Schema Course sem AggregateRating (oportunidade de estrelas no Google)',
]:
    p = doc.add_paragraph(style='List Bullet')
    if 'CRITICO' in item:
        run = p.add_run(item)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 66, 28)
    else:
        p.add_run(item)

# ── SAVE ──
path_local = 'D:/EspanholComVoce/site-vendas/AUDITORIA_FINAL_V2.docx'
doc.save(path_local)
print(f'Saved: {path_local}')

path_drive = 'G:/Meu Drive/EspanholComVoce/AUDITORIA_FINAL_V2.docx'
os.makedirs(os.path.dirname(path_drive), exist_ok=True)
shutil.copy2(path_local, path_drive)
print(f'Copied: {path_drive}')
