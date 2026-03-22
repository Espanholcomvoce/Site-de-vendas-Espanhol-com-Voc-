"""
Gera o documento Word completo do plano de lançamento de 7 dias
Espanhol com Você — @espanholcomvoce
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ─── Estilos ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(22)
style_h1.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
style_h1.font.bold = True

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(16)
style_h2.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
style_h2.font.bold = True

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(13)
style_h3.font.color.rgb = RGBColor(0xb4, 0x82, 0x00)
style_h3.font.bold = True


def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 60)
    run.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)
    run.font.size = Pt(8)


def add_box(doc, title, content, color=RGBColor(0x0a, 0x16, 0x28)):
    p = doc.add_paragraph()
    run = p.add_run(f'▌ {title}')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    p2 = doc.add_paragraph(content)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.style.font.size = Pt(10)


def add_story(doc, num, horario, formato, visual, texto_tela, ale_fala=None, instrucao=None, nota=None):
    p = doc.add_paragraph()
    run = p.add_run(f'STORY {num} — {horario}')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)

    info = doc.add_paragraph()
    run_f = info.add_run(f'Formato: ')
    run_f.font.bold = True
    run_f.font.size = Pt(10)
    info.add_run(formato).font.size = Pt(10)

    if visual:
        info2 = doc.add_paragraph()
        run_v = info2.add_run(f'Visual: ')
        run_v.font.bold = True
        run_v.font.size = Pt(10)
        info2.add_run(visual).font.size = Pt(10)

    t = doc.add_paragraph()
    run_t = t.add_run('Texto na tela:')
    run_t.font.bold = True
    run_t.font.size = Pt(10)
    for line in texto_tela.split('\n'):
        p_line = doc.add_paragraph(line)
        p_line.paragraph_format.left_indent = Cm(1)
        p_line.style.font.size = Pt(10)

    if ale_fala:
        a = doc.add_paragraph()
        run_a = a.add_run('Ale fala:')
        run_a.font.bold = True
        run_a.font.size = Pt(10)
        run_a.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
        for line in ale_fala.split('\n'):
            p_line = doc.add_paragraph(line)
            p_line.paragraph_format.left_indent = Cm(1)
            run_l = p_line.runs[0] if p_line.runs else p_line.add_run(line)
            run_l.font.size = Pt(10)
            run_l.font.italic = True

    if instrucao:
        i = doc.add_paragraph()
        run_i = i.add_run('Instrução de gravação: ')
        run_i.font.bold = True
        run_i.font.size = Pt(9)
        run_i.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
        i.add_run(instrucao).font.size = Pt(9)

    if nota:
        n = doc.add_paragraph()
        run_n = n.add_run('⚠ Nota: ')
        run_n.font.bold = True
        run_n.font.size = Pt(9)
        n.add_run(nota).font.size = Pt(9)

    doc.add_paragraph()  # espaço


def add_slide(doc, num, tipo_slide, conteudo):
    p = doc.add_paragraph()
    run = p.add_run(f'SLIDE {num} — {tipo_slide}')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x38, 0xbd, 0xf8)

    for line in conteudo.split('\n'):
        if line.strip():
            pl = doc.add_paragraph(line)
            pl.paragraph_format.left_indent = Cm(1)
            pl.style.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ESPANHOL COM VOCÊ')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('@espanholcomvoce')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PLANO DE LANÇAMENTO — 7 DIAS')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Operação: De zero vendas orgânicas para as primeiras 5-10')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Roteiros de Feed • Roteiros de Stories • Emails de Lançamento')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Março 2026')
run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════

doc.add_heading('ÍNDICE', level=1)
doc.add_paragraph()

indice = [
    'PARTE 1 — VISÃO GERAL DA OPERAÇÃO',
    '  Estrutura dos 7 dias',
    '  Checklist de pré-lançamento',
    '  Regras gerais de gravação',
    '',
    'PARTE 2 — ROTEIROS DE STORIES (83 Stories)',
    '  Dia 1 — Ativação (9 Stories)',
    '  Dia 2 — Dor (11 Stories)',
    '  Dia 3 — Método (11 Stories)',
    '  Dia 4 — Prova Social (11 Stories)',
    '  Dia 5 — Mini-Aula Gratuita (17 Stories)',
    '  Dia 6 — Oferta (12 Stories)',
    '  Dia 7 — Fechamento (12 Stories)',
    '',
    'PARTE 3 — ROTEIROS DE FEED (7 Publicações)',
    '  Dia 1 — Reel Viral: Expressões',
    '  Dia 2 — Carrossel: Dor e Identificação',
    '  Dia 3 — Reel: Quebra de Crenças',
    '  Dia 4 — Carrossel: Prova Social',
    '  Dia 5 — Reel: Mini-Aula Pública',
    '  Dia 6 — Carrossel: Oferta Direta',
    '  Dia 7 — Reel: Fechamento',
    '',
    'PARTE 4 — EMAILS DE LANÇAMENTO (5 Emails)',
    '  Email 1 (Dia 4) — Aquecimento',
    '  Email 2 (Dia 5) — Antecipação',
    '  Email 3 (Dia 6) — Abertura da Oferta',
    '  Email 4 (Dia 7 manhã) — Último Dia',
    '  Email 5 (Dia 7 noite) — Última Chance',
    '',
    'PARTE 5 — PROJEÇÃO DE VENDAS E DMs',
]

for item in indice:
    if not item.strip():
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(item)
    if p.runs:
        if item.startswith('PARTE'):
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(11)
        else:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PARTE 1 — VISÃO GERAL
# ═══════════════════════════════════════════════════════════

doc.add_heading('PARTE 1 — VISÃO GERAL DA OPERAÇÃO', level=1)
add_divider(doc)

doc.add_heading('Conceito', level=2)
doc.add_paragraph(
    'Ale abre uma "Semana de Imersão" — 7 dias de conteúdo intenso gratuito nos Stories e Feed, '
    'culminando numa oferta especial com prazo real. É a primeira vez que o orgânico será usado '
    'como canal de venda estruturado. Meta: 5-10 vendas orgânicas em 7 dias, sem investir em anúncios.'
)

doc.add_heading('Estrutura dos 7 Dias', level=2)

table = doc.add_table(rows=8, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Grid Accent 1'

headers = ['Dia', 'Tema', 'Feed', 'Stories', 'Venda?']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

data = [
    ['1 (Segunda)', 'ATIVAÇÃO', 'Reel viral', '9 stories', 'NÃO'],
    ['2 (Terça)', 'DOR', 'Carrossel dor', '11 stories', 'NÃO'],
    ['3 (Quarta)', 'MÉTODO', 'Reel crenças', '11 stories', 'NÃO'],
    ['4 (Quinta)', 'PROVA SOCIAL', 'Carrossel prova', '11 stories', 'NÃO'],
    ['5 (Sexta)', 'MINI-AULA', 'Reel mini-aula', '17 stories', 'NÃO'],
    ['6 (Sábado)', 'OFERTA', 'Carrossel venda', '12 stories', 'SIM ✓'],
    ['7 (Domingo)', 'FECHAMENTO', 'Reel fechamento', '12 stories', 'SIM ✓'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('Checklist de Pré-Lançamento', level=2)
doc.add_paragraph('Fazer ANTES do Dia 1:')

checklist = [
    '☐ TESTAR MANYCHAT — Comentar TRAVA, METODO, AULA, FALSOS, QUERO num post e verificar se o bot responde e entrega a isca',
    '☐ TESTAR SEQUÊNCIA DE EMAIL — Cadastrar email teste no Mailchimp via ManyChat e verificar se os emails disparam',
    '☐ TESTAR VSL — Assistir a VSL inteira, verificar se carrega e se o botão de compra funciona',
    '☐ ATUALIZAR LINK NA BIO — CTA claro → Linktree/página com: Aula gratuita (VSL) + PDF gratuito + WhatsApp',
    '☐ ESCOLHER O BÔNUS DA SEMANA — Algo que custe zero/pouco para entregar mas tenha valor percebido alto (recomendação: aula ao vivo com Ale em grupo pequeno)',
    '☐ SEPARAR 4 DEPOIMENTOS FORTES — 1 por avatar: viajante, profissional, acadêmico, morar fora (prints WhatsApp, vídeos, mensagens reais)',
    '☐ PREPARAR OS 5 EMAILS de lançamento (Parte 4 deste documento)',
    '☐ CONFIGURAR KEYWORD "QUERO" no ManyChat — deve enviar link direto para página de vendas',
]
for item in checklist:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.5)
    p.runs[0].font.size = Pt(10)

doc.add_heading('Regras Gerais de Gravação', level=2)

regras = [
    'Tom: Ale falando direto pra câmera como amiga. Natural, sem teleprompter, sem produção pesada.',
    'Visual: Fundo natural (casa, escritório, sofá). Ring light ou janela. Nada de cenário montado.',
    'Texto na tela: Sempre colocar legenda resumida + stickers de engajamento.',
    'Duração Stories: 10-15 segundos cada. Máximo 20.',
    'Duração Reels: 30-60 segundos. Gancho nos primeiros 3 segundos.',
    'SEMPRE: "Nós, nativos, falamos..." — primeira pessoa.',
    'NUNCA: "Os nativos falam..." — terceira pessoa proibida.',
    'NUNCA: "Fluente em 30 dias", "Fácil", "Decore" — contra a marca.',
    'Carrosséis: Identidade v2 — navy (#0a1628), dourado (#fbbf24), Inter 900, Caveat manuscrito.',
    'Reels: Câmera frontal, vertical 9:16, enquadramento peito pra cima.',
]
for r in regras:
    p = doc.add_paragraph(r, style='List Bullet')
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PARTE 2 — STORIES
# ═══════════════════════════════════════════════════════════

doc.add_heading('PARTE 2 — ROTEIROS DE STORIES', level=1)
p = doc.add_paragraph('83 Stories em 7 dias — roteiro completo com texto, falas e instruções')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
add_divider(doc)

# ── DIA 1 ──
doc.add_heading('DIA 1 — ATIVAÇÃO (Acordar a audiência)', level=2)
add_box(doc, 'OBJETIVO', 'Fazer os seguidores interagirem. Treinar o algoritmo. ZERO venda.')

add_story(doc, 1, '09:00', 'Enquete', 'Fundo navy (#0a1628) + texto dourado (#fbbf24)',
    'Responde com sinceridade 👇\n\nENQUETE:\n"Você entende espanhol mas trava na hora de falar?"\n[ SIM 😩 ]  [ NÃO 😎 ]',
    nota='Essa enquete é a base de tudo. Resposta esperada: 80%+ SIM.')

add_story(doc, 2, '09:05', 'Vídeo de Ale (10 seg)', 'Close, olhando pra câmera',
    '"Fica comigo essa semana 👀"',
    '"Se você respondeu SIM... fica comigo essa semana. Eu tenho uma coisa pra te mostrar que vai mudar a forma como você aprende espanhol."')

add_story(doc, 3, '12:00', 'Caixinha de perguntas', 'Foto da Ale sorrindo + caixinha',
    'Eu quero te ajudar DE VERDADE\n\nCAIXINHA:\n"Qual sua maior dificuldade com espanhol? Me conta aqui 👇"',
    nota='Ale deve ler TODAS as respostas e escolher 3-4 para responder em vídeo.')

add_story(doc, 4, '15:00', 'Vídeo de Ale respondendo pergunta 1 (15 seg)', 'Close, tom didático',
    'Pergunta do seguidor no topo + resposta em destaque',
    '"Recebi essa pergunta aqui: [lê a pergunta]\n\nOlha, isso é mais comum do que você imagina. [Resposta curta]\n\nNo espanhol real, nós falamos assim: [ensina]"')

add_story(doc, 5, '15:05', 'Vídeo de Ale respondendo pergunta 2 (15 seg)', 'Mesmo formato',
    'Pergunta + resposta',
    '"Essa outra pergunta aqui eu AMEI: [lê a pergunta]\n\nSabe por que tanta gente tem essa dúvida? Porque na escola ensinam errado. A gente fala assim: [ensina]"')

add_story(doc, 6, '15:10', 'Vídeo de Ale respondendo pergunta 3 (15 seg)', 'Mesmo formato',
    'Pergunta + resposta + "Manda mais perguntas 👆"',
    '"Última por enquanto: [lê pergunta]\n\nEssa é clássica. Brasileiro SEMPRE erra isso. [Resposta]\n\nManda mais perguntas que eu respondo amanhã também."')

add_story(doc, 7, '18:00', 'Texto + emoji slider', 'Fundo navy + texto dourado',
    'Essa semana eu vou fazer algo\nque nunca fiz aqui.\n\nPrepara. 🔥\n\nEMOJI SLIDER: 🔥 "Quão curiosa você tá?"')

add_story(doc, 8, '20:00', 'Vídeo de Ale (15 seg)', 'Ale relaxada, iluminação quente',
    '"[X]% travam 😱" + "Amanhã eu explico por quê"',
    '"Gente, vocês viram o resultado da enquete de hoje? [X]% de vocês travam na hora de falar espanhol. Eu sei EXATAMENTE por que isso acontece. Porque eu passei pela MESMA coisa com o português.\n\nAmanhã eu te conto essa história. Não perde."')

add_story(doc, 9, '20:05', 'Enquete de fechamento', 'Fundo navy',
    'Quer que eu conte minha história\nde quando cheguei no Brasil amanhã?\n\nENQUETE:\n[ SIM, conta! ]  [ Já conheço ]',
    nota='Quem responder "Já conheço" provavelmente é seguidor antigo e mais quente.')

p = doc.add_paragraph()
run = p.add_run('TOTAL DIA 1: 9 Stories | Interações: 2 enquetes + 1 caixinha + 1 slider')
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_page_break()

# ── DIA 2 ──
doc.add_heading('DIA 2 — DOR (Fazer a pessoa se reconhecer)', level=2)
add_box(doc, 'OBJETIVO', 'Ativar a dor emocional. Fazer o seguidor pensar "sou eu". Conexão com a história da Ale.')

add_story(doc, 1, '09:00', 'Texto estático', 'Fundo preto puro + texto branco, fonte grande',
    '2012.\n\nEu cheguei no Brasil\ncom um certificado avançado\nde português.\n\nAchava que estava preparada.',
    nota='Sem stickers, sem nada. Só o texto. Silêncio visual. Cria tensão.')

add_story(doc, 2, '09:01', 'Texto estático', 'Fundo preto + texto branco',
    'Fui ao mercado.\n\nO cara me perguntou:\n"Quer sacola?"\n\nEu entendi.\nMas na hora de responder...\n\nnada saiu.')

add_story(doc, 3, '09:02', 'Vídeo de Ale (20 seg)', 'Close, tom emocional mas controlado',
    '"Eu entendia tudo. Mas não falava nada."',
    '"Gente, eu tinha estudado português por ANOS. Eu tinha certificado. AVANÇADO. E quando cheguei aqui... eu travei. No mercado. Com o porteiro. Com a corretora. Eu entendia TUDO. Mas não saía NADA.\n\nSabe aquela sensação? De ter as palavras na cabeça mas elas não chegam na boca?\n\nFoi a pior sensação que eu já tive."',
    instrucao='Não precisa chorar, mas precisa ser real. Olhar direto na câmera.')

add_story(doc, 4, '09:03', 'Enquete', 'Fundo navy + dourado',
    'Isso já aconteceu com VOCÊ\ntentando falar espanhol?\n\nENQUETE:\n[ SIM, igualzinho 😩 ]  [ Tenho medo que aconteça ]')

add_story(doc, 5, '12:00', 'Vídeo de Ale (15 seg)', 'Ale mais leve, tom de descoberta',
    '"O problema não era eu. Era o método."',
    '"Depois daquele dia no mercado, eu percebi uma coisa: o problema NÃO era eu. Era o MÉTODO que eu tinha usado pra aprender.\n\nEu tinha decorado gramática. Tinha feito exercício. Tinha tirado nota boa.\n\nMas ninguém me ensinou a FALAR."')

add_story(doc, 6, '12:01', 'Texto estático', 'Fundo navy + texto dourado grande',
    'A escola ensina você\na passar na prova.\n\nNão a falar\ncom um ser humano.')

add_story(doc, 7, '15:00', 'Vídeo de Ale (15 seg)', 'Ale mais animada, energia subindo',
    '"Mergulhei no idioma real."',
    '"Aí eu comecei a estudar português de um jeito COMPLETAMENTE diferente. Sem decorar. Sem traduzir. Mergulhei no idioma REAL. No que as pessoas falavam na rua, na TV, no trabalho.\n\nE em pouco tempo... destravei. Consegui emprego numa multinacional. Em português."')

add_story(doc, 8, '15:01', 'Texto estático', 'Fundo navy',
    'Esse é o mesmo método\nque eu uso pra ensinar VOCÊ.\n\nMais de 5.000 alunos\njá destravaram com ele.\n\nEle tem nome:\nMétodo Imersão Nativa®')

add_story(doc, 9, '18:00', 'Vídeo de Ale (10 seg)', 'Ale direta, olho na câmera',
    '"Amanhã: por que seu cérebro trava 🧠"',
    '"Se você se identificou com a minha história... se você trava... se você entende mas não fala...\n\nFica comigo essa semana. Sério.\n\nAmanhã eu vou te mostrar POR QUE o seu cérebro trava. E como consertar."')

add_story(doc, 10, '20:00', 'Resultado da enquete + vídeo (10 seg)', 'Print do resultado',
    '"[X]% se identificaram"',
    '"Olha isso. [X]% de vocês se identificaram com a minha história. Vocês não estão sozinhos. E eu vou ajudar cada um. Amanhã, 9 da manhã, aqui."')

add_story(doc, 11, '20:05', 'Caixinha de perguntas', 'Fundo navy',
    'Me conta: em que situação\nvocê mais trava em espanhol?\n\nCAIXINHA:\n"Descreve a situação aqui 👇"',
    nota='Respostas alimentam conteúdo dos próximos dias e identificam avatares.')

p = doc.add_paragraph()
run = p.add_run('TOTAL DIA 2: 11 Stories | Arco: Tensão → Identificação → Esperança → Curiosidade')
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_page_break()

# ── DIA 3 ──
doc.add_heading('DIA 3 — MÉTODO (Mostrar que a solução existe)', level=2)
add_box(doc, 'OBJETIVO', 'Explicar o método. Atacar objeção "não sei o que vou receber". Mostrar plataforma por dentro.')

add_story(doc, 1, '09:00', 'Texto estático provocativo', 'Fundo preto + texto branco grande',
    'Quanto mais você ESTUDA espanhol...\n\npior você fala.\n\nSim, é sério.')

add_story(doc, 2, '09:01', 'Vídeo de Ale (20 seg)', 'Tom de "vou te contar um segredo"',
    '"Ouve → Traduz → Pensa → Traduz → Fala = TRAVA"',
    '"Parece loucura, né? Mas é assim que funciona.\n\nQuando você estuda espanhol do jeito tradicional, seu cérebro cria um hábito: TRADUZIR.\n\nAlguém fala com você em espanhol. Seu cérebro ouve em espanhol, traduz pro português, pensa a resposta em português, traduz pro espanhol... e quando você vai abrir a boca? A pessoa já foi embora."')

add_story(doc, 3, '09:02', 'Vídeo de Ale (15 seg)', None,
    '"Pensar direto em espanhol = fluência"',
    '"O método que funciona é o oposto. É PARAR de traduzir. É treinar seu cérebro pra pensar direto em espanhol.\n\nE sabe como se faz isso? IMERSÃO. Contato real com o idioma real. Do jeito que nós, nativos, aprendemos."')

add_story(doc, 4, '12:00', 'Screen recording da plataforma', 'Ale gravando tela do celular mostrando as aulas',
    '"180+ aulas de 5 minutos"',
    '"Olha como funciona por dentro. Essa aqui é uma aula real do programa.\n\n[Mostra lista de aulas]\nSão mais de 180 aulas. Cada uma tem 5 minutos.\n\n[Clica numa aula e dá play por 5 seg]\nOlha: eu falo em espanhol, tem legenda em espanhol e em português.\n\n[Mostra material de apoio]\nE aqui tem o material escrito, exercício, áudio..."',
    nota='ESSE É UM DOS STORIES MAIS IMPORTANTES. Resolve objeção "não sei o que vou receber".')

add_story(doc, 5, '12:01', 'Screen recording continuação', None,
    '"Imersão Ativa = espanhol real todo dia"',
    '"E aqui é a parte que faz a diferença: a Imersão Ativa.\n\n[Mostra seção de imersão]\nTodo dia você recebe contato com espanhol real. Vídeos, áudios, textos — do jeito que nativos falam.\n\nNão é exercício chato. É espanhol de verdade, do dia a dia."')

add_story(doc, 6, '12:02', 'Vídeo de Ale (10 seg)', None,
    '"Suporte humano por WhatsApp em até 24h"',
    '"E a parte que meus alunos mais amam: o suporte por WhatsApp.\n\nQualquer dúvida, você manda mensagem e em até 24 horas minha equipe responde. Personalizado. Não é robô."')

add_story(doc, 7, '15:00', 'Enquete', 'Fundo navy + dourado',
    'Responde sincera:\n\nENQUETE:\n"Você consegue dedicar 5 minutos por dia pra aprender espanhol?"\n[ SIM, consigo ✅ ]  [ Nem isso 😅 ]',
    nota='Esperado: 85-90% SIM. Planta semente de que o método cabe na rotina.')

add_story(doc, 8, '15:05', 'Vídeo de Ale (10 seg)', 'Ale apontando pro resultado',
    '"5 minutos por dia = 1 aula completa"',
    '"[X]% de vocês têm 5 minutos por dia.\n\nSe você tem 5 minutos, eu consigo te ensinar. Sério. É o tempo de UMA aula do meu programa.\n\n5 minutos. Não é desculpa."')

add_story(doc, 9, '18:00', 'Texto estático', 'Fundo navy + texto dourado',
    '3 pilares.\n180+ aulas de 5 minutos.\nImersão com espanhol real.\nSuporte humano por WhatsApp.\nAcesso pra sempre.\n\nE funciona pra qualquer nível.')

add_story(doc, 10, '20:00', 'Vídeo de Ale (15 seg)', 'Ale casual, vibe noturna',
    '"Amanhã: resultados reais 🔥"',
    '"Amanhã eu vou trazer os resultados. Alunos REAIS. Pessoas como você que travavam e hoje falam espanhol.\n\nNão é depoimento falso. São prints de WhatsApp, vídeos, histórias reais.\n\nAmanhã. Não perde."')

add_story(doc, 11, '20:05', 'Enquete de fechamento', None,
    'Depois de ver como o método funciona por dentro, o que você achou?\n\nENQUETE:\n[ Quero saber mais 🔥 ]  [ Não é pra mim ]',
    nota='Quem responder "Quero saber mais" entra na lista de DMs do Dia 6.')

p = doc.add_paragraph()
run = p.add_run('TOTAL DIA 3: 11 Stories | Objeções atacadas: "não sei o que vou receber" + "não tenho tempo"')
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_page_break()

# ── DIA 4 ──
doc.add_heading('DIA 4 — PROVA SOCIAL (Resultados reais)', level=2)
add_box(doc, 'OBJETIVO', 'Eliminar dúvidas com depoimentos reais. Preparar terreno emocional para mini-aula.')

add_story(doc, 1, '09:00', 'Texto estático', 'Fundo navy + texto branco',
    '"Ale, mas será que funciona PRA MIM?"\n\nEssa é a pergunta que eu mais recebo.\n\nA resposta está nos próximos Stories.')

add_story(doc, 2, '09:01', 'Print WhatsApp — Avatar VIAJANTE', 'Screenshot real + fundo navy',
    'ELA VIAJOU PRA ARGENTINA E SE VIROU SOZINHA 🇦🇷\n\n[Print da mensagem da aluna]\n\n"Antes do programa, ela dependia do Google Tradutor. Hoje conversa com nativos."')

add_story(doc, 3, '09:02', 'Print WhatsApp — Avatar PROFISSIONAL', 'Mesmo formato',
    'ELE FECHOU NEGÓCIO EM ESPANHOL 💼\n\n[Print da mensagem do aluno]\n\n"Antes, ele travava em reuniões com clientes do Chile. Hoje apresenta em espanhol sem medo."')

add_story(doc, 4, '12:00', 'Print WhatsApp — Avatar ACADÊMICO', 'Mesmo formato',
    'ELA PASSOU NO DELE 🎓\n\n[Print da mensagem da aluna]\n\n"Precisava da certificação pra bolsa de estudos. Passou na primeira tentativa."')

add_story(doc, 5, '12:01', 'Print WhatsApp — Avatar MORAR FORA', 'Mesmo formato',
    'ELA SE MUDOU PRA ESPANHA E SE INTEGROU 🇪🇸\n\n[Print da mensagem da aluna]\n\n"Chegou sem falar nada. Hoje trabalha e vive em espanhol."')

add_story(doc, 6, '12:02', 'Vídeo de Ale (15 seg)', 'Ale emocionada',
    '"5.000+ alunos reais"',
    '"Gente, eu tenho MILHARES de mensagens assim. Cada uma dessas histórias é uma pessoa real que travava — igualzinho a mim quando cheguei aqui.\n\nJá são mais de 5.000 alunos. Cada um com a sua história."')

add_story(doc, 7, '15:00', 'Vídeo de aluno ou print adicional', None,
    'ANTES: "eu travava em tudo"\nDEPOIS: [vídeo/print com resultado]\n\n"X meses de programa"',
    nota='Se não tiver vídeo de aluno, usar print de WhatsApp com detalhes específicos.')

add_story(doc, 8, '15:01', 'Enquete / Quiz', 'Fundo navy',
    'Qual dessas histórias mais se parece com o SEU objetivo?\n\nENQUETE:\n[ Viajar sem travar ]  [ Trabalho/carreira ]\n\n(ou Quiz com 4 opções: Viagem / Morar fora / Estudo / Trabalho)',
    nota='Essa resposta identifica o avatar. Usar no Dia 6 para personalizar DMs.')

add_story(doc, 9, '18:00', 'Texto estático', 'Fundo navy + dourado grande',
    'O próximo resultado real\npode ser o SEU.\n\nAmanhã eu vou te dar\numa MINI-AULA gratuita\naqui nos Stories.\n\nA melhor aula que eu já dei de graça.')

add_story(doc, 10, '20:00', 'Vídeo de Ale (15 seg)', 'Ale empolgada, energia alta',
    '"Amanhã 9h — Mini-aula gratuita 🔔"',
    '"Amanhã é o dia mais importante dessa semana. Eu vou dar uma mini-aula REAL aqui nos Stories. Vou te ensinar 3 expressões que nós, nativos, usamos todo dia — e que você não aprendeu na escola.\n\nE no final eu tenho uma surpresa. Ativa o sininho pra não perder. 9 da manhã."')

add_story(doc, 11, '20:05', 'Countdown sticker', 'Fundo navy + countdown',
    'MINI-AULA GRATUITA\namanhã 9h\n\n[STICKER COUNTDOWN: "Mini-aula da Ale" — amanhã 09:00]',
    nota='O sticker de countdown envia notificação para quem ativar.')

p = doc.add_paragraph()
run = p.add_run('TOTAL DIA 4: 11 Stories | Depoimentos: 4 (1 por avatar)')
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_page_break()

# ── DIA 5 ──
doc.add_heading('DIA 5 — MINI-AULA GRATUITA (Ponto de virada)', level=2)
add_box(doc, 'OBJETIVO', 'Entregar valor absurdo de graça. Mostrar método funcionando em tempo real. Criar desejo. Antecipar oferta.')
add_box(doc, 'IMPORTANTE', 'Sequência mais longa e mais importante da semana. Ale precisa gravar com energia, clareza e emoção.', RGBColor(0xfb, 0xbf, 0x24))

add_story(doc, 1, '09:00', 'Vídeo de Ale (10 seg)', 'Ale pronta, sorrindo, energia alta',
    '"MINI-AULA GRATUITA 🎓"',
    '"Bom dia! Hoje é dia de mini-aula. Se você quer aprender espanhol DE VERDADE, fica comigo nos próximos Stories. Eu vou te ensinar o que ninguém te ensinou na escola."')

add_story(doc, 2, '09:01', 'Texto estático', 'Fundo preto + texto branco',
    'Antes da aula, uma pergunta:\n\nVocê sabe a diferença entre\nESTUDAR um idioma\ne ADQUIRIR um idioma?\n\nA maioria não sabe.\nE é por isso que trava.')

add_story(doc, 3, '09:02', 'Vídeo de Ale (15 seg)', None,
    '"Estudar ≠ Adquirir"',
    '"ESTUDAR é decorar regra, fazer exercício, traduzir texto. Isso treina a parte ERRADA do cérebro.\n\nADQUIRIR é fazer o que uma criança faz: ouvir, repetir, errar, tentar de novo, dentro de um CONTEXTO real.\n\nNenhum nativo aprendeu espanhol estudando gramática. A gente aprendeu VIVENDO.\n\nO meu método reproduz isso pra você."')

add_story(doc, 4, '09:03', 'Texto de transição', 'Fundo navy + texto dourado',
    'Agora eu vou te ensinar\n3 expressões que nós, nativos,\nusamos TODO DIA.\n\nCoisas que a escola nunca te ensinou.\n\nRepete comigo. Em voz alta. Sério.')

add_story(doc, 5, '09:04', 'Vídeo — EXPRESSÃO 1 (15 seg)', 'Ale + pílula visual',
    'Pílula 1: "COMO FALAR" (dourado)\nPílula 2: "TÔ NEM AÍ" (branco/preto)\nPílula 3: "EM ESPANHOL" (dourado)',
    '"Expressão 1: quando você quer dizer \'tô nem aí\' em espanhol.\n\nNa escola te ensinam: \'no me importa\'. Mas nós, nativos, falamos: \'me da igual\'.\n\nRepete: ME DA IGUAL.\n\n[Pausa]\n\n\'¿Que van a decir de ti? Me da igual.\' Perfeito."')

add_story(doc, 6, '09:05', 'Vídeo — EXPRESSÃO 2 (15 seg)', 'Mesmo formato',
    'Pílula: "COMO FALAR" / "FAZER DE CONTA" / "EM ESPANHOL"',
    '"Expressão 2: quando você quer dizer \'fazer de conta\'.\n\nMuito brasileiro fala \'hacer de cuenta\'. ERRADO. Nós falamos: \'HACERSE EL/LA...\'\n\nPor exemplo: \'Se hace el tonto.\' — Ele faz de desentendido.\n\nRepete: SE HACE EL TONTO."')

add_story(doc, 7, '09:06', 'Vídeo — EXPRESSÃO 3 (15 seg)', 'Mesmo formato',
    'Pílula: "COMO FALAR" / "DAR UMA OLHADINHA" / "EM ESPANHOL"',
    '"Expressão 3: quando você quer dizer \'dar uma olhadinha\'.\n\nNão existe tradução direta. Nós falamos: \'ECHAR UN VISTAZO\'.\n\n\'Voy a echar un vistazo.\' — Vou dar uma olhadinha.\n\nRepete: ECHAR UN VISTAZO. Essa é das minhas favoritas."')

add_story(doc, 8, '09:07', 'Enquete', 'Fundo navy',
    'E aí? Conseguiu?\n\nENQUETE:\n"Qual expressão você mais curtiu?"\n[ Me da igual 😎 ]  [ Echar un vistazo 👀 ]')

add_story(doc, 9, '09:08', 'Vídeo de Ale (15 seg) — Transição', None,
    '"Imagina isso TODO DIA"',
    '"Viu? 2 minutos e você aprendeu 3 coisas que nunca te ensinaram na escola.\n\nAgora imagina isso TODO DIA. 5 minutos por dia. 180 aulas assim. Com suporte pra tirar dúvida. Com imersão em espanhol real. Com acesso pra SEMPRE."')

add_story(doc, 10, '09:09', 'Vídeo de Ale (10 seg) — Antecipação', None,
    '"Amanhã: condição especial 🔥"',
    '"Isso é o Programa Imersão Nativa. É o método que eu criei baseado em como EU aprendi português.\n\nE amanhã eu vou abrir uma condição ESPECIAL pra quem quer começar. Algo que eu nunca fiz antes. Fica de olho."')

add_story(doc, 11, '09:10', 'Texto + caixinha', 'Fundo navy',
    'Se você tem alguma dúvida sobre\no método ou sobre como funciona,\nme pergunta aqui 👇\n\nCAIXINHA: "Sua dúvida sobre o programa"',
    nota='As dúvidas viram conteúdo de quebra de objeção nos Stories do Dia 6.')

add_story(doc, '12-14', '14:00', 'Vídeo de Ale respondendo dúvidas (3 Stories)', None,
    'Respostas às dúvidas recebidas',
    'Story 12 — Dúvida sobre TEMPO:\n"Cada aula tem 5 minutos. CINCO. Você assiste no ônibus, na fila do banco, antes de dormir. Não precisa de hora marcada. É no seu ritmo."\n\nStory 13 — Dúvida sobre NÍVEL:\n"Funciona pra qualquer nível. Do zero ao avançado. As aulas são progressivas. Você começa de onde faz sentido pra você."\n\nStory 14 — Dúvida sobre PREÇO/GARANTIA:\n"Tem garantia de 7 dias. Se você não gostar, eu devolvo todo o dinheiro. Sem pergunta. E dá pra parcelar em 12x."')

add_story(doc, 15, '18:00', 'Print de depoimento + texto', None,
    'Ela começou com a mesma dúvida que você tem agora.\n\n[Print de aluna que tinha a mesma objeção e superou]\n\n"X meses depois..."')

add_story(doc, 16, '20:00', 'Vídeo de Ale (15 seg) — Grande antecipação', None,
    '"AMANHÃ — Condição exclusiva ⏰"',
    '"Gente. Amanhã é o dia.\n\nEu vou abrir o Programa Imersão Nativa com um bônus EXCLUSIVO que eu nunca ofereci.\n\n[Descreve o bônus em 1 frase]\n\nSó vale amanhã e domingo. Depois não volta.\n\nAmanhã cedo eu te conto tudo."')

add_story(doc, 17, '20:05', 'Countdown sticker', 'Fundo navy + countdown',
    'OFERTA EXCLUSIVA\namanhã 9h\n\n[STICKER COUNTDOWN: "Abertura Imersão Nativa" — amanhã 09:00]')

p = doc.add_paragraph()
run = p.add_run('TOTAL DIA 5: 17 Stories | Arco: Ensina → Encanta → Mostra possibilidade → Antecipa oferta')
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_page_break()

# ── DIA 6 ──
doc.add_heading('DIA 6 — OFERTA (Abrir a venda)', level=2)
add_box(doc, 'OBJETIVO', 'VENDER. Clareza total na oferta. Quebrar objeções uma a uma. Links em todos os Stories. DMs personalizadas.')
add_box(doc, 'DMs', 'Enviar 50-100 DMs personalizadas ao longo do dia para quem interagiu nos Stories da semana.', RGBColor(0xfb, 0xbf, 0x24))

add_story(doc, 1, '09:00', 'Vídeo de Ale (15 seg)', 'Ale séria mas acolhedora',
    '"ABRIU 🔓"',
    '"Bom dia. Essa semana eu te contei minha história, te mostrei o método por dentro, trouxe resultados de alunos reais, e te dei uma mini-aula gratuita.\n\nAgora é a hora de você decidir.\n\nO Programa Imersão Nativa está aberto. E eu tenho uma condição especial que só vale HOJE e AMANHÃ."')

add_story(doc, 2, '09:01', 'Texto estático — O QUE INCLUI', 'Fundo navy + texto branco organizado',
    'PROGRAMA IMERSÃO NATIVA®\n\n✅ 180+ aulas em vídeo de 5 minutos\n✅ Legendas em espanhol e português\n✅ Material de apoio bilíngue\n✅ Imersão Ativa com espanhol real\n✅ Suporte por WhatsApp em até 24h\n✅ Acesso VITALÍCIO\n✅ Garantia de 7 dias')

add_story(doc, 3, '09:02', 'Texto estático — BÔNUS', 'Fundo dourado (#fbbf24) + texto navy — DESTAQUE',
    '🎁 BÔNUS EXCLUSIVO\nSÓ ESTE FIM DE SEMANA:\n\n[Nome do bônus]\n[Descrição em 1-2 linhas]\n\n⏰ Válido até AMANHÃ (domingo) 23:59\nDepois não volta.')

add_story(doc, 4, '09:03', 'Texto estático — PREÇO', 'Fundo navy + preço em dourado grande',
    'De R$597\nPor R$397\nou 12x de R$39,70\n\nMenos que um café por dia.\nAcesso pra SEMPRE.\nGarantia de 7 dias.')

add_story(doc, 5, '09:04', 'Vídeo de Ale + LINK (10 seg)', 'Ale apontando pra cima (pro link)',
    '"LINK AQUI 👆" + adesivo de link para página de vendas',
    '"Se você quer começar, o link tá aqui em cima. É seguro, é pela Hotmart, e tem garantia de 7 dias. Se em 7 dias você não gostar, eu devolvo TUDO. Sem pergunta nenhuma."')

add_story(doc, 6, '11:00', 'Vídeo — Objeção #1: TEMPO (15 seg)', None,
    '"Aulas de 5 minutos. Sem desculpa." + link',
    '"\'Ale, eu não tenho tempo.\'\n\nEu entendo. Eu também sou ocupada. Mas as aulas têm CINCO MINUTOS. Cinco. Menos que um TikTok.\n\nSe você tem 5 minutos por dia, você tem tempo pra aprender espanhol."')

add_story(doc, 7, '13:00', 'Print de notificação Hotmart ou texto', 'Screenshot + celebração',
    '🔥 VENDA #[X]!\n\nMais uma pessoa destravando o espanhol HOJE.\n\n"Bem-vinda, [nome]! Te vejo nas aulas 💛"',
    nota='Se tiver venda real, mostrar. Se não tiver ainda, pular e postar quando tiver a primeira.')

add_story(doc, 8, '14:00', 'Vídeo — Objeção #2: JÁ TENTEI (15 seg)', None,
    '"O problema não é você. É o método." + link',
    '"\'Ale, eu já tentei aprender espanhol antes e não consegui.\'\n\nEU TAMBÉM. Eu estudei português na escola e não consegui falar quando cheguei no Brasil.\n\nSabe por quê? O método era errado. Quando eu mudei o método, tudo mudou. É exatamente isso que eu ofereço pra você."')

add_story(doc, 9, '16:00', 'Vídeo — Objeção #3: PREÇO (15 seg)', None,
    '"12x R$39,70 — acesso vitalício" + link',
    '"\'Ale, tá caro.\'\n\nEu entendo a preocupação. Mas faz a conta: R$39,70 por mês. Parcelado em 12x. Acesso pra SEMPRE. Não é mensalidade.\n\nQuanto custa PERDER aquela viagem por não falar? PERDER aquela promoção? PERDER aquela bolsa?\n\n39 reais por mês. Pra mudar sua vida."')

add_story(doc, 10, '17:00', 'Depoimento forte + link', 'Print impactante',
    '"Eu também achava que não ia conseguir. Olha agora."\n\n[Print de depoimento com resultado específico]\n\n👆 Quer ser o próximo? Link aqui em cima.\n\n+ link para página de vendas')

add_story(doc, 11, '19:00', 'Vídeo — Urgência (10 seg)', None,
    '"⏰ Fecha amanhã 23:59" + link',
    '"Gente, o bônus [nome do bônus] fecha AMANHÃ à meia-noite. Não vai ter extensão, não vai ter \'último dia de verdade\'. Amanhã 23:59 e acabou.\n\nSe você quer, é agora."')

add_story(doc, 12, '21:00', 'Enquete de fechamento', None,
    'Quero entender:\n\nENQUETE:\n"O que te impede de começar?"\n[ Preciso pensar mais ]  [ Vou comprar! ]',
    nota='Quem responder "Preciso pensar mais" recebe DM no Dia 7: "Oi! Vi que você tá pensando. O que te deixou na dúvida? Posso ajudar. 💛"')

p = doc.add_paragraph()
run = p.add_run('TOTAL DIA 6: 12 Stories | Links para página: mínimo 5 | DMs: 50-100 mensagens')
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_page_break()

# ── DIA 7 ──
doc.add_heading('DIA 7 — FECHAMENTO (Últimas horas)', level=2)
add_box(doc, 'OBJETIVO', 'Urgência máxima. Último dia. Converter indecisos. Fechar com emoção e gratidão.')

add_story(doc, 1, '08:00', 'Texto estático', 'Fundo navy + texto branco grande',
    'ÚLTIMO DIA.\n\nO bônus [nome] sai hoje 23:59.\n\nNão vai voltar.')

add_story(doc, 2, '09:00', 'Vídeo de Ale (15 seg)', 'Ale calma, tom de conversa final',
    '"Hoje é o dia." + link',
    '"Bom dia. Último dia.\n\nEu não vou ficar te pressionando. Você sabe se precisa ou não aprender espanhol. Você sabe se o método faz sentido pra você.\n\nEu só quero que você tome a decisão HOJE. Porque amanhã o bônus sai e a vida continua igual."')

add_story(doc, 3, '10:00', 'Compilação de prints', '4-6 prints pequenos em colagem',
    'O que meus alunos falam:\n\n[Colagem com frases destacadas:\n"mudou minha vida"\n"deveria ter começado antes"\n"melhor investimento"\n"as aulas são viciantes"]')

add_story(doc, 4, '12:00', 'Vídeo — A pergunta que decide (15 seg)', None,
    '"Daqui a 6 meses, o que muda?" + link',
    '"Eu quero te fazer uma pergunta. Sem julgamento.\n\nDaqui a 6 meses... se você não fizer NADA diferente... você vai falar espanhol?\n\n[Pausa 2 seg]\n\nNão. Vai continuar travando. Vai continuar traduzindo na cabeça. Vai continuar dizendo \'um dia eu aprendo\'.\n\nOu você pode começar HOJE."',
    instrucao='Pausas longas. Na pausa de 2 seg, olhar pro lado como se pensasse junto com a pessoa. Tom empático, não acusatório.')

add_story(doc, 5, '14:00', 'Texto — Resumo final + link', 'Fundo navy, organizado',
    'ÚLTIMAS HORAS ⏰\n\n✅ 180+ aulas de 5 minutos\n✅ Acesso vitalício\n✅ Suporte WhatsApp 24h\n✅ Garantia 7 dias\n🎁 Bônus: [nome] (SÓ HOJE)\n\n12x R$39,70\n\n👆 Link aqui em cima')

add_story(doc, 6, '16:00', 'Vídeo — Depoimento ao vivo (15 seg)', None,
    '"Aluna nova — menos de 24h" + link',
    '"Olha essa mensagem que eu recebi AGORA de alguém que comprou ontem:\n\n[Mostra print ou lê]\n\n\'Ale, já assisti 3 aulas e tô chocada. É tudo que eu precisava.\'\n\nIsso me emociona de verdade. É pra isso que eu faço o que faço."')

add_story(doc, 7, '18:00', 'Texto — Contagem regressiva + link', 'Fundo navy + dourado GRANDE',
    '6 HORAS.\n\nDepois disso, o bônus [nome]\nnão volta.\n\nVocê decide.')

add_story(doc, 8, '20:00', 'Vídeo de Ale (15 seg) — Emocional', 'Ale íntima, sentada no sofá',
    '"Obrigada 💛" + link',
    '"Gente, eu quero agradecer todo mundo que me acompanhou essa semana.\n\nQuem comprou: te vejo nas aulas. Você tomou a melhor decisão.\n\nQuem ainda tá na dúvida: você tem até meia-noite. Depois eu volto pro conteúdo normal e essa condição fecha."')

add_story(doc, 9, '21:00', 'Texto + link', 'Fundo preto + texto branco',
    '3 horas.\n\nLink aqui em cima 👆')

add_story(doc, 10, '22:00', 'Texto + link', 'Fundo preto + dourado',
    '2 horas.')

add_story(doc, 11, '23:00', 'Vídeo rápido (10 seg) + link', None,
    '"ÚLTIMA HORA ⏰" + link',
    '"Última hora. Se você quer, é agora. Link aqui em cima. Meia-noite fecha."')

add_story(doc, 12, '23:59', 'Vídeo de fechamento', 'Ale relaxada, sorrindo',
    '"Obrigada 💛 Nos vemos nas aulas."',
    '"Fechou.\n\nObrigada a cada pessoa que confiou em mim essa semana.\n\nPra quem comprou: amanhã eu mando uma mensagem especial pra vocês no WhatsApp.\n\nPra quem não comprou: tudo bem. Eu continuo aqui, ensinando espanhol de graça todo dia. Quando for a hora, você sabe onde me encontrar. 💛\n\nBoa noite."')

p = doc.add_paragraph()
run = p.add_run('TOTAL DIA 7: 12 Stories | Links em TODOS (exceto último) | DMs: recontactar todos que não compraram')
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PARTE 3 — FEED
# ═══════════════════════════════════════════════════════════

doc.add_heading('PARTE 3 — ROTEIROS DE FEED', level=1)
p = doc.add_paragraph('7 publicações — 4 Reels + 3 Carrosséis — roteiro pronto para gravar')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
add_divider(doc)

# ── FEED DIA 1 ──
doc.add_heading('DIA 1 — REEL VIRAL: "Dando uma olhadinha"', level=2)
add_box(doc, 'DADOS', 'Tipo: Expressões Idiomáticas (TOPO) | Horário: 19:00 | Duração: 30-40 seg | ManyChat: NÃO')

doc.add_heading('Roteiro segundo a segundo', level=3)

sections_d1 = [
    ('GANCHO (0-3s)', 'Texto na tela: "Você fala isso ERRADO em espanhol 👇"\n\nAle fala: "Aposto que você fala isso errado em espanhol."\n\nInstrução: Olhar direto na câmera. Expressão provocativa, não arrogante. Tom de amiga que vai te corrigir.'),
    ('CONTEXTO (3-8s)', 'Pílulas visuais:\n• "COMO FALAR" (fundo dourado)\n• "DANDO UMA OLHADINHA" (fundo branco, texto preto bold)\n• "EM ESPANHOL" (fundo dourado)\n\nAle fala: "Quando você vai numa loja e quer falar \'tô só dando uma olhadinha\'... o que você fala em espanhol?"\n\nInstrução: Tom curioso. Pausa de 1 seg depois da pergunta.'),
    ('ERRO COMUM (8-15s)', 'Texto na tela: ❌ "Estoy solo mirando"\n\nAle fala: "A maioria fala \'estoy solo mirando\'. Dá pra entender? Dá. Mas nenhum nativo fala assim."\n\nInstrução: Balançar cabeça sutilmente no "nenhum nativo".'),
    ('EXPRESSÃO REAL (15-25s)', 'Texto na tela: ✅ "Solo estoy echando un vistazo"\n\nAle fala: "Nós, nativos, falamos: \'Solo estoy echando un vistazo.\'\n\n[Devagar] \'E-chan-do un vis-ta-zo.\'\n[Normal] \'Solo estoy echando un vistazo.\'\n\nRepete comigo. Vai."\n\nInstrução: Sorriso na expressão certa. Devagar 1ª vez, normal 2ª. Pausa 1.5 seg pro "Vai".'),
    ('EXEMPLO (25-32s)', 'Texto na tela: "No puedo irme sin echar un vistazo" 👀\n\nAle fala: "Imagina que você tá numa loja na Espanha e a vendedora pergunta se precisa de ajuda:\n\n[Muda o tom] \'¿Necesita algo?\'\n\'No, gracias. Solo estoy echando un vistazo.\'\n\nPerfeito. Natural. Como nativa."'),
    ('CTA (32-38s)', 'Texto na tela: "SALVA 📌 e manda pra alguém que precisa"\n\nAle fala: "Salva esse vídeo porque você vai precisar. E manda pra aquela pessoa que sempre fala espanhol errado. Ela precisa ver isso."\n\nInstrução: Tom leve, brincalhão. Apontar pra baixo.'),
]
for title, content in sections_d1:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x38, 0xbd, 0xf8)
    for line in content.split('\n'):
        pl = doc.add_paragraph(line)
        pl.paragraph_format.left_indent = Cm(0.5)
        pl.style.font.size = Pt(10)

doc.add_heading('Legenda', level=3)
doc.add_paragraph(
    'Você fala "dando uma olhadinha" em espanhol ERRADO. 👀\n\n'
    'A maioria dos brasileiros fala "estoy solo mirando". Dá pra entender? Dá. Mas nenhum nativo fala assim.\n\n'
    'Nós falamos: "Solo estoy echando un vistazo."\n\n'
    'Parece diferente? É porque É diferente. Na escola te ensinam o espanhol do livro. Aqui eu te ensino o espanhol REAL.\n\n'
    'Salva esse post e manda pra alguém que precisa 📌\n\n'
    '#espanholcomvoce #aprendaespanhol #espanholreal #expressoes #comofalarem #espanholnativo '
    '#espanholparabrasileiros #espanholonline #espanhol #dicasdeespanhol #espanholparaviagem '
    '#aprendendo #idiomasnaweb #espanholdodia #linguaespanhola'
).style.font.size = Pt(10)

doc.add_page_break()

# ── FEED DIA 2 ──
doc.add_heading('DIA 2 — CARROSSEL: "5 sinais de que você aprende errado"', level=2)
add_box(doc, 'DADOS', 'Tipo: Dor e Identificação (MEIO) | Horário: 19:00 | 7 slides | ManyChat: TRAVA → 20 Erros')

slides_d2 = [
    ('Slide 1 — CAPA IMPACTO', 'Fundo navy + glow dourado | Foto Ale expressão séria\nHeadline: "5 sinais de que você está aprendendo espanhol ERRADO"\nDestaque: "ERRADO" em dourado | Manuscrito: "será que é o seu caso? →"'),
    ('Slide 2 — TEXTO PURO', 'Navy sólido | "01" dourado grande\n"Você traduz TUDO na cabeça antes de abrir a boca"\nSubtexto: "Ouve → traduz → pensa → traduz → e a pessoa já foi embora."'),
    ('Slide 3 — CENA REAL', 'Foto Ale + tag pílula "SINAL 02"\n"Você entende filmes mas TRAVA numa conversa real"\nManuscrito: "assistir não é falar..."'),
    ('Slide 4 — TEXTO PURO', '"03" dourado\n"Você estuda há MESES mas não monta uma frase sozinho"\nSubtexto: "Sabe as palavras mas não sabe como juntar."'),
    ('Slide 5 — SPLIT CARD', 'Dois cards lado a lado:\nCard esquerdo: "SINAL 04" — "Você sabe gramática mas não consegue USAR"\nCard direito: "SINAL 05" — "Você tem VERGONHA de falar na frente de nativos"'),
    ('Slide 6 — CARD PAPEL', 'Card creme sobre fundo escuro | Tag: "A VERDADE"\n"Se você marcou 3 ou mais... O problema NÃO é você. É o método."\nManuscrito: "eu sei porque eu vivi isso →"'),
    ('Slide 7 — CTA', 'Foto Ale sorrindo | Keyword "TRAVA" outline dourado\n"Se você se identificou, comenta TRAVA que eu te mando um diagnóstico gratuito dos 20 erros mais comuns 🎁"\nManuscrito: "te mando na hora 🚀"'),
]
for title, content in slides_d2:
    add_slide(doc, title.split(' — ')[0].replace('Slide ', ''), title.split(' — ')[1] if ' — ' in title else '', content)

doc.add_heading('Legenda', level=3)
doc.add_paragraph(
    'Você está aprendendo espanhol errado. E nem sabe. 😶\n\n'
    'Eu digo isso porque eu VIVI isso. Quando cheguei no Brasil em 2012, eu tinha certificado AVANÇADO de português. E no primeiro dia eu travei.\n\n'
    'Se você: → Traduz tudo na cabeça → Entende mas não fala → Estuda há meses e não monta frase → Sabe gramática mas não usa → Tem vergonha de falar\n\n'
    'Eu sei exatamente o que você tá passando. E eu sei como resolver.\n\n'
    'Comenta TRAVA que eu te mando um diagnóstico gratuito dos 20 erros mais comuns 🎁\n\n'
    '#espanholcomvoce #aprendaespanhol #travanoespanhol #identificacao #voceseidentifica '
    '#espanholreal #espanholparabrasileiros #espanhol #aprenderespanhol #espanholonline '
    '#dicasdeespanhol #erroscomuns #metodo #imersaonativa #espanholdodia'
).style.font.size = Pt(10)

doc.add_page_break()

# ── FEED DIA 3 ──
doc.add_heading('DIA 3 — REEL: "Quanto mais estuda, pior fala"', level=2)
add_box(doc, 'DADOS', 'Tipo: Quebra de Crenças (MEIO) | Horário: 19:00 | Duração: 35-45 seg | ManyChat: METODO → Guia Pronúncia')

sections_d3 = [
    ('GANCHO (0-3s)', 'Texto: "Quanto mais você ESTUDA espanhol... PIOR você fala."\n\nAle fala: "Quanto mais você estuda espanhol... pior você fala."\n\nInstrução: Expressão de verdade dura. Pausa 1.5 seg. Tom sério, não agressivo.'),
    ('EXPLICAÇÃO (3-20s)', 'Texto: "Seu cérebro foi treinado para TRADUZIR"\n\nAle fala: "Parece loucura, né? Mas pensa comigo. Quando você estuda do jeito tradicional, seu cérebro cria um vício: TRADUZIR. Ouve em espanhol, traduz pro português, pensa a resposta, traduz de volta... e quando vai abrir a boca...\n\n[Pausa dramática]\n\n...a pessoa já mudou de assunto."\n\nInstrução: Contar nos dedos cada etapa. Ritmo cada vez mais rápido.'),
    ('SOLUÇÃO (20-35s)', 'Texto: "IMERSÃO > Tradução"\n\nAle fala: "O caminho é o oposto. Parar de traduzir. Parar de decorar. Mergulhar no espanhol REAL.\n\nÉ assim que nós, nativos, aprendemos. Nenhuma criança aprende a falar estudando gramática.\n\nÉ por imersão. Por prática. Por contato REAL com o idioma. E é exatamente isso que o meu método faz."\n\nInstrução: Tom muda de sério para confiante. Sorrir em "nós, nativos".'),
    ('CTA (35-42s)', 'Texto: "Comenta METODO 👇"\n\nAle fala: "Se você quer entender como funciona, comenta METODO que eu te mando um guia gratuito de como nativos realmente aprendem idiomas."\n\nInstrução: Tom de convite. Apontar pra baixo. Sorriso final.'),
]
for title, content in sections_d3:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x38, 0xbd, 0xf8)
    for line in content.split('\n'):
        pl = doc.add_paragraph(line)
        pl.paragraph_format.left_indent = Cm(0.5)
        pl.style.font.size = Pt(10)

doc.add_page_break()

# ── FEED DIA 4 ──
doc.add_heading('DIA 4 — CARROSSEL: "Eles travavam. Hoje falam."', level=2)
add_box(doc, 'DADOS', 'Tipo: Prova Social (FUNDO) | Horário: 19:00 | 8 slides | ManyChat: AULA → VSL')

slides_d4 = [
    ('1', 'CAPA IMPACTO', 'Navy + glow dourado | Ale sorrindo\n"Eles travavam. Hoje falam espanhol."\nDestaque: "falam espanhol" dourado | Manuscrito: "resultados reais →"'),
    ('2', 'CARD PAPEL — Viajante', 'Tag: "RESULTADO REAL" | Print WhatsApp\n"O VIAJANTE 🌎"\nAntes: "Dependia do Google Tradutor"\nDepois: "Pedi comida, negociei preço, fiz amigos"'),
    ('3', 'CARD PAPEL — Profissional', 'Tag: "RESULTADO REAL" | Print WhatsApp\n"O PROFISSIONAL 💼"\nAntes: "Travava em reuniões"\nDepois: "Apresentei pro diretor de Santiago"'),
    ('4', 'CARD PAPEL — Acadêmico', 'Tag: "RESULTADO REAL" | Print WhatsApp\n"A ACADÊMICA 🎓"\nAntes: "Medo de reprovar"\nDepois: "Aprovada no DELE de primeira"'),
    ('5', 'CARD PAPEL — Morar Fora', 'Tag: "RESULTADO REAL" | Print WhatsApp\n"A QUE FOI MORAR FORA 🏠"\nAntes: "Cheguei perdida"\nDepois: "Consegui emprego em espanhol"'),
    ('6', 'TEXTO PURO — Números', 'Navy | "5.000+" dourado gigante\n"alunos já destravaram com o Método Imersão Nativa®"\nManuscrito: "e o próximo pode ser você"'),
    ('7', 'CENA REAL — Ale', 'Foto Ale feliz | Tag: "POR QUE FUNCIONA"\n"Eu criei o método baseado em como EU aprendi português. Por imersão. Por prática."\nManuscrito: "eu vivi isso na pele"'),
    ('8', 'CTA', 'Ale confiante | Keyword "AULA" outline dourado\n"Quer ver como funciona? Comenta AULA que eu te mando uma aula demonstrativa gratuita 🎁"\nManuscrito: "te mando agora 🚀"'),
]
for num, tipo, content in slides_d4:
    add_slide(doc, num, tipo, content)

doc.add_page_break()

# ── FEED DIA 5 ──
doc.add_heading('DIA 5 — REEL: "5 frases em 60 segundos"', level=2)
add_box(doc, 'DADOS', 'Tipo: Expressões + Método (TOPO/MEIO) | Horário: 12:00 | Duração: 45-60 seg | ManyChat: AULA → VSL')

sections_d5 = [
    ('GANCHO (0-3s)', 'Texto: "5 frases em espanhol em 60 SEGUNDOS ⏱️"\n\nAle fala: "Em 60 segundos você vai aprender 5 frases que nós, nativos, usamos TODO DIA."\n\nInstrução: Energia alta. Mostrar 5 com a mão. Tom de desafio.'),
    ('FRASE 1 (3-12s)', 'Texto: 🇧🇷 "Tô nem aí" → 🇪🇸 "Me da igual"\n\nAle: "1. \'Tô nem aí\': ME DA IGUAL. \'¿Qué piensan? Me da igual.\' Repete!"'),
    ('FRASE 2 (12-22s)', 'Texto: 🇧🇷 "Dar uma olhadinha" → 🇪🇸 "Echar un vistazo"\n\nAle: "2. \'Dar uma olhadinha\': ECHAR UN VISTAZO. \'Solo estoy echando un vistazo.\' Perfeito!"'),
    ('FRASE 3 (22-32s)', 'Texto: 🇧🇷 "Cara de pau" → 🇪🇸 "Caradura"\n\nAle: "3. \'Cara de pau\': CARADURA. Uma palavra só. \'¡Qué caradura!\' Fácil, né?"'),
    ('FRASE 4 (32-42s)', 'Texto: 🇧🇷 "Sei lá" → 🇪🇸 "Qué sé yo"\n\nAle: "4. \'Sei lá\': QUÉ SÉ YO. Bem informal. \'¿Dónde vamos? Qué sé yo.\' Amo essa!"'),
    ('FRASE 5 (42-50s)', 'Texto: 🇧🇷 "Ficar de bobeira" → 🇪🇸 "Estar al pedo"\n\nAle: "5. \'Ficar de bobeira\': ESTAR AL PEDO. Sim, a gente fala isso! [Ri] \'Estaba al pedo en casa.\' Na escola NUNCA te ensinam essa."'),
    ('CTA (50-60s)', 'Texto: "Quer MAIS? Comenta AULA 👇"\n\nAle: "5 frases em menos de 1 minuto. Imagina o que você aprende com o método completo. 180 aulas assim. De 5 minutos.\n\nComenta AULA que eu te mando uma aula demonstrativa gratuita. Salva esse vídeo. Você vai precisar."'),
]
for title, content in sections_d5:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x38, 0xbd, 0xf8)
    for line in content.split('\n'):
        pl = doc.add_paragraph(line)
        pl.paragraph_format.left_indent = Cm(0.5)
        pl.style.font.size = Pt(10)

doc.add_page_break()

# ── FEED DIA 6 ──
doc.add_heading('DIA 6 — CARROSSEL DE VENDA: "Eu nunca fiz isso antes"', level=2)
add_box(doc, 'DADOS', 'Tipo: CTA Direto (FUNDO) | Horário: 12:00 | 8 slides | ManyChat: QUERO → Página de vendas')

slides_d6 = [
    ('1', 'CAPA IMPACTO', 'Navy + glow dourado forte | Ale olhando pra câmera\n"Eu nunca fiz isso antes."\nManuscrito: "e não vou fazer de novo tão cedo →"'),
    ('2', 'TEXTO PURO', '"Essa semana eu te contei minha história. Te mostrei o método por dentro. Trouxe resultados de alunos reais. Te dei uma mini-aula gratuita.\n\nAgora eu quero te fazer um convite."\nManuscrito: "um convite de verdade"'),
    ('3', 'CENA REAL — O programa', 'Foto Ale ensinando | Tag: "O PROGRAMA"\n"Programa Imersão Nativa®"\n✅ 180+ aulas de 5 minutos\n✅ Legendas espanhol + português\n✅ Imersão com espanhol real\n✅ Suporte WhatsApp 24h\n✅ Acesso VITALÍCIO'),
    ('4', 'SPLIT CARD — Antes vs Depois', 'Card ✗ ANTES: "Traduz tudo / Trava / Vergonha / Não progride"\nCard ✓ DEPOIS: "Pensa em espanhol / Fala com confiança / Conversa com nativos / Progresso real"'),
    ('5', 'CARD PAPEL — Bônus', 'Tag: "🎁 BÔNUS EXCLUSIVO" | Card creme\n"SÓ ESTE FIM DE SEMANA\n[Nome do bônus]\n[Descrição]\n⏰ Até domingo 23:59"'),
    ('6', 'TEXTO PURO — Preço', 'Navy | Riscado: "R$597" | Dourado gigante: "R$397"\n"ou 12x de R$39,70"\nManuscrito: "menos que um café por dia"\nSubtexto: "Acesso vitalício. Garantia 7 dias."'),
    ('7', 'CARD PAPEL — Garantia', 'Tag: "GARANTIA ZERO RISCO" | Card creme\n"Se em 7 dias você não gostar, eu devolvo 100% do seu dinheiro.\nSem pergunta. Sem burocracia. Sem letras miúdas.\nO risco é todo meu."\nManuscrito: "eu confio no método"'),
    ('8', 'CTA FINAL', 'Ale sorrindo, braços abertos | Keyword "QUERO" gigante dourado\n"Comenta QUERO e eu te mando o link direto pra começar agora"\nManuscrito: "te vejo nas aulas 💛"\nSubtexto: "ou clica no link da bio 👆"'),
]
for num, tipo, content in slides_d6:
    add_slide(doc, num, tipo, content)

doc.add_page_break()

# ── FEED DIA 7 ──
doc.add_heading('DIA 7 — REEL FECHAMENTO: "A última vez"', level=2)
add_box(doc, 'DADOS', 'Tipo: CTA Direto (FUNDO) | Horário: 12:00 | Duração: 25-35 seg | ManyChat: QUERO → Página de vendas')

sections_d7 = [
    ('GANCHO (0-3s)', 'Texto: "Essa é a última vez que eu falo sobre isso."\n\nAle fala a mesma frase.\n\nInstrução: Tom sério, não triste. Olho na câmera. Pausa 1 seg.'),
    ('REFLEXÃO (3-18s)', 'Texto: "A pergunta que decide tudo"\n\nAle: "Eu quero te fazer uma pergunta. Daqui a 6 meses... se você não fizer NADA diferente... você vai falar espanhol?\n\n[Pausa 2 seg]\n\nNão. Vai continuar travando. Vai continuar traduzindo na cabeça. Vai continuar dizendo \'um dia eu aprendo\'.\n\nOu você pode começar HOJE."\n\nInstrução: Pausas longas. Olhar pro lado na pausa de 2 seg. Tom empático, não acusatório. Energia sobe no "HOJE".'),
    ('OFERTA (18-28s)', 'Texto: "12x R$39,70 | Acesso vitalício | Garantia 7 dias"\n\nAle: "Programa Imersão Nativa. 180 aulas. 5 minutos cada. Suporte por WhatsApp. Acesso pra sempre. 12x de 39 reais e 70.\n\nE hoje é o ÚLTIMO dia com o bônus [nome]. Meia-noite fecha. De verdade."\n\nInstrução: Clareza nos detalhes. Enfatizar "de verdade".'),
    ('CTA (28-35s)', 'Texto: "Comenta QUERO 👇 ou link na bio 👆"\n\nAle: "Comenta QUERO que eu te mando o link. Ou clica na bio.\n\nSe for pra você, você sabe. Te vejo nas aulas."\n\n[Ale sorri e acena]\n\nInstrução: Tom de despedida carinhosa. Acenar como "tchau, te espero lá".'),
]
for title, content in sections_d7:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x38, 0xbd, 0xf8)
    for line in content.split('\n'):
        pl = doc.add_paragraph(line)
        pl.paragraph_format.left_indent = Cm(0.5)
        pl.style.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PARTE 4 — EMAILS
# ═══════════════════════════════════════════════════════════

doc.add_heading('PARTE 4 — EMAILS DE LANÇAMENTO', level=1)
p = doc.add_paragraph('5 emails — disparar via Mailchimp para toda a base de leads')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
add_divider(doc)

# EMAIL 1
doc.add_heading('EMAIL 1 — Aquecimento (Dia 4, manhã)', level=2)

p = doc.add_paragraph()
run = p.add_run('Assunto: ')
run.font.bold = True
p.add_run('"Eu preciso te contar uma coisa"')

p = doc.add_paragraph()
run = p.add_run('Preview text: ')
run.font.bold = True
p.add_run('"Você já consegue falar espanhol com confiança?"')

doc.add_paragraph()
email1 = """Oi [nome],

Aqui é a Ale, do Espanhol com Você.

Você baixou o [nome da isca] há um tempo. Eu quero te fazer uma pergunta sincera:

Você já consegue falar espanhol com confiança?

Se a resposta for não, eu entendo. Porque eu passei pela mesma coisa.

Quando cheguei no Brasil em 2012, eu tinha certificado avançado de português. AVANÇADO. E no primeiro dia eu travei. No mercado, com o porteiro, com a corretora. Eu entendia tudo — mas não conseguia falar nada.

Sabe o que eu descobri? O problema não era eu. Era o método.

Essa semana eu estou fazendo algo diferente no meu Instagram. Estou compartilhando tudo o que eu sei sobre por que as pessoas travam — e como resolver isso.

Mas eu tenho uma coisa especial preparada pra você. Algo que eu nunca fiz antes.

Fica de olho no seu email amanhã. Sério.

Um abraço,
Ale

P.S. Se você quer acompanhar o que estou fazendo essa semana, me segue no Instagram: @espanholcomvoce"""

p = doc.add_paragraph(email1)
p.style.font.size = Pt(10)

doc.add_page_break()

# EMAIL 2
doc.add_heading('EMAIL 2 — Antecipação (Dia 5, manhã)', level=2)

p = doc.add_paragraph()
run = p.add_run('Assunto: ')
run.font.bold = True
p.add_run('"Amanhã tem algo especial (e não vai durar)"')

p = doc.add_paragraph()
run = p.add_run('Preview text: ')
run.font.bold = True
p.add_run('"Uma condição que eu nunca ofereci antes"')

doc.add_paragraph()
email2 = """Oi [nome],

Essa semana eu fiz algo que nunca tinha feito:

→ Contei minha história de quando travei no Brasil
→ Mostrei o Método Imersão Nativa por dentro
→ Trouxe resultados reais de alunos
→ Dei uma mini-aula gratuita nos Stories

Tudo isso pra te mostrar que falar espanhol de verdade é possível. Mesmo que você já tenha tentado antes e não conseguido.

Agora vem a parte especial:

AMANHÃ eu vou abrir o Programa Imersão Nativa com um bônus que eu nunca ofereci antes:

🎁 [NOME DO BÔNUS]
[Descrição em 1 linha]

Esse bônus só vai estar disponível amanhã (sábado) e domingo. Depois fecha. De verdade.

O programa:
✅ 180+ aulas de 5 minutos — no seu ritmo
✅ Imersão com espanhol real de nativos
✅ Suporte por WhatsApp em até 24h
✅ Acesso VITALÍCIO
✅ Garantia de 7 dias

Tudo isso por 12x de R$39,70 — menos que um café por dia.

Amanhã cedo eu mando o link. Fica de olho.

Um abraço,
Ale

P.S. Se em 7 dias você não gostar, eu devolvo 100% do seu dinheiro. Sem pergunta. O risco é todo meu."""

p = doc.add_paragraph(email2)
p.style.font.size = Pt(10)

doc.add_page_break()

# EMAIL 3
doc.add_heading('EMAIL 3 — Abertura da Oferta (Dia 6, 09:00)', level=2)

p = doc.add_paragraph()
run = p.add_run('Assunto: ')
run.font.bold = True
p.add_run('"Abriu — e o bônus é esse"')

p = doc.add_paragraph()
run = p.add_run('Preview text: ')
run.font.bold = True
p.add_run('"Programa Imersão Nativa com condição exclusiva"')

doc.add_paragraph()
email3 = """Oi [nome],

Chegou a hora.

O Programa Imersão Nativa está aberto com uma condição que eu nunca ofereci — e que só vale até amanhã (domingo) às 23:59.

━━━━━━━━━━━━━━━━━━━━

O QUE VOCÊ RECEBE:

✅ 180+ aulas em vídeo de 5 minutos
   Todas em espanhol com legendas em espanhol e português

✅ Imersão Ativa
   Contato diário com espanhol real — vídeos, áudios, textos nativos

✅ Suporte por WhatsApp em até 24h
   Dúvidas respondidas de forma personalizada pela equipe pedagógica

✅ Material de apoio bilíngue
   Exercícios interativos e áudios de nativos

✅ Acesso VITALÍCIO
   Você paga uma vez e acessa pra sempre

✅ Garantia incondicional de 7 dias
   Não gostou? Devolvo 100% do seu dinheiro. Sem pergunta.

━━━━━━━━━━━━━━━━━━━━

🎁 BÔNUS EXCLUSIVO (SÓ ESTE FIM DE SEMANA):

[NOME DO BÔNUS]
[Descrição em 2-3 linhas]

⏰ Disponível até domingo 23:59. Depois não volta.

━━━━━━━━━━━━━━━━━━━━

INVESTIMENTO:

De R$597 por R$397
Ou 12x de R$39,70

Menos que um café por dia. Acesso pra sempre.

━━━━━━━━━━━━━━━━━━━━

👉 [BOTÃO: QUERO COMEÇAR AGORA]

━━━━━━━━━━━━━━━━━━━━

[nome], eu criei esse método porque eu vivi o que você tá vivendo. Eu travei. Eu tive vergonha. Eu achei que o problema era eu.

Não era. Era o método.

Mais de 5.000 alunos já provaram que funciona. O próximo pode ser você.

Te vejo nas aulas,
Ale 💛

P.S. Lembra: garantia de 7 dias. Se você não gostar, não precisa explicar — eu devolvo tudo. O risco é zero pra você."""

p = doc.add_paragraph(email3)
p.style.font.size = Pt(10)

doc.add_page_break()

# EMAIL 4
doc.add_heading('EMAIL 4 — Último Dia (Dia 7, 09:00)', level=2)

p = doc.add_paragraph()
run = p.add_run('Assunto: ')
run.font.bold = True
p.add_run('"Último dia — hoje 23:59"')

p = doc.add_paragraph()
run = p.add_run('Preview text: ')
run.font.bold = True
p.add_run('"Depois disso o bônus não volta"')

doc.add_paragraph()
email4 = """Oi [nome],

Direto ao ponto: hoje é o último dia.

O Programa Imersão Nativa com o bônus [nome do bônus] fecha hoje à meia-noite. Depois disso, o bônus sai e não volta.

Eu não vou ficar te pressionando. Você sabe se precisa ou não aprender espanhol.

Mas eu quero te fazer uma pergunta:

Daqui a 6 meses, se você não mudar nada... você vai falar espanhol?

Se a resposta é não, talvez hoje seja o dia de mudar isso.

✅ 180+ aulas de 5 minutos
✅ Acesso vitalício
✅ Suporte WhatsApp 24h
✅ Garantia 7 dias
🎁 Bônus: [nome] — SÓ HOJE

12x de R$39,70

👉 [BOTÃO: QUERO COMEÇAR AGORA]

Te vejo nas aulas,
Ale 💛"""

p = doc.add_paragraph(email4)
p.style.font.size = Pt(10)

doc.add_page_break()

# EMAIL 5
doc.add_heading('EMAIL 5 — Última Chance (Dia 7, 20:00)', level=2)

p = doc.add_paragraph()
run = p.add_run('Assunto: ')
run.font.bold = True
p.add_run('"3 horas"')

p = doc.add_paragraph()
run = p.add_run('Preview text: ')
run.font.bold = True
p.add_run('"Meia-noite fecha"')

doc.add_paragraph()
email5 = """[nome],

3 horas.

Meia-noite o bônus [nome] sai do Programa Imersão Nativa.

Se você quer destravar seu espanhol de uma vez, é agora.

180+ aulas. 5 minutos cada. Acesso pra sempre. Garantia de 7 dias.

12x de R$39,70.

👉 [BOTÃO: QUERO COMEÇAR AGORA]

Depois de meia-noite, eu volto pro conteúdo normal e essa condição fecha.

Quando for a hora, você sabe onde me encontrar. 💛

Ale"""

p = doc.add_paragraph(email5)
p.style.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PARTE 5 — PROJEÇÃO E DMs
# ═══════════════════════════════════════════════════════════

doc.add_heading('PARTE 5 — PROJEÇÃO DE VENDAS E ESTRATÉGIA DE DMs', level=1)
add_divider(doc)

doc.add_heading('Projeção de vendas por canal', level=2)

table2 = doc.add_table(rows=6, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Light Grid Accent 1'

headers2 = ['Canal', 'Pessimista', 'Realista', 'Otimista']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)

data2 = [
    ['Stories → Link → Compra', '1', '4', '9'],
    ['ManyChat → VSL → Compra', '1', '2', '4'],
    ['Email blast → VSL → Compra', '0', '1', '2'],
    ['DMs manuais → Compra', '1', '3', '6'],
    ['TOTAL VENDAS', '3', '10', '21'],
]
for row_idx, row_data in enumerate(data2):
    for col_idx, val in enumerate(row_data):
        cell = table2.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if row_idx == 4:
                    r.font.bold = True

doc.add_paragraph()

doc.add_heading('Estratégia de DMs — O canal #1 da semana', level=2)

doc.add_paragraph(
    'As DMs são o canal que mais vai converter nesta semana. '
    'É manual e não escalável — mas para sair do zero, contato direto com quem engajou é o caminho mais curto.'
)

doc.add_heading('Dia 6 — Primeiro round de DMs (50-100 mensagens)', level=3)
doc.add_paragraph('Quem contatar:')
targets = [
    'Todos que responderam enquetes nos Stories durante a semana',
    'Todos que comentaram TRAVA, METODO ou AULA no feed',
    'Todos que mandaram perguntas na caixinha',
    'Todos que responderam "Quero saber mais" na enquete do Dia 3',
]
for t in targets:
    doc.add_paragraph(t, style='List Bullet').runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_heading('Modelo de DM (personalizar, NÃO copiar/colar)', level=3)
doc.add_paragraph(
    '"Oi [nome]! Vi que você respondeu [X] nos meus Stories essa semana. '
    'Queria te dizer que hoje eu abri uma condição especial pro Programa '
    'Imersão Nativa com [bônus]. Se tiver interesse, posso te mandar o link. 💛"'
).runs[0].font.italic = True

doc.add_paragraph()
doc.add_heading('Dia 7 — Segundo round (recontato)', level=3)
doc.add_paragraph('Quem recontatar:')
targets2 = [
    'Quem abriu DM ontem mas não comprou',
    'Quem disse "vou ver" ou "depois"',
    'Quem engajou nos Stories do Dia 7',
    'Quem respondeu "Preciso pensar mais" na enquete do Dia 6',
]
for t in targets2:
    doc.add_paragraph(t, style='List Bullet').runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_heading('Modelo de recontato', level=3)
doc.add_paragraph(
    '"Oi [nome], só passando pra avisar que hoje é o último dia com o bônus. '
    'Se precisar de ajuda pra decidir, me fala. 💛"'
).runs[0].font.italic = True

doc.add_paragraph()
add_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nDocumento gerado em Março 2026\n')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Espanhol com Você — @espanholcomvoce\n')
run.font.size = Pt(12)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plano de Lançamento Orgânico — Operação 7 Dias')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# ── Salvar ──
output_path = r'D:\EspanholComVoce\lancamento_7dias.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
print(f'Tamanho: {os.path.getsize(output_path) / 1024:.0f} KB')
