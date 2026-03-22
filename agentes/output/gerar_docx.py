"""
Converts calendario_abril_roteiros.md to a formatted Word document.
Processes the markdown line by line for efficiency.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# Colors
NAVY = RGBColor(0x0a, 0x16, 0x28)
GOLD = RGBColor(0xb4, 0x82, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

INPUT_FILE = r"D:\EspanholComVoce\agentes\output\calendario_abril_roteiros.md"
OUTPUT_FILE = r"D:\EspanholComVoce\agentes\output\calendario_abril_roteiros.docx"


def set_run_font(run, size=11, bold=False, color=DARK_GRAY, font_name="Calibri"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    # Set East Asian and Complex Script fonts too
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)


def add_styled_paragraph(doc, text, size=11, bold=False, color=DARK_GRAY, alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if alignment is not None:
        pf.alignment = alignment
    return p


def add_cover_page(doc):
    # Add some blank lines for vertical centering
    for _ in range(6):
        add_styled_paragraph(doc, "", size=11)

    add_styled_paragraph(doc, "CALENDARIO ABRIL 2026", size=32, bold=True,
                         color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "ROTEIROS COMPLETOS", size=28, bold=True,
                         color=GOLD, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_styled_paragraph(doc, "Espanhol com Voce — @espanholcomvoce", size=16, bold=True,
                         color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_styled_paragraph(doc, "90 posts: 30 Reels + 60 Carrosseis", size=14, bold=False,
                         color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_styled_paragraph(doc, "Marco 2026", size=12, bold=False,
                         color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Page break
    doc.add_page_break()


def add_toc_page(doc, days):
    add_styled_paragraph(doc, "SUMARIO", size=22, bold=True, color=NAVY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_styled_paragraph(doc, "", size=6)

    for day_info in days:
        add_styled_paragraph(doc, day_info, size=11, bold=False, color=DARK_GRAY, space_after=4)

    doc.add_page_break()


def process_inline_formatting(paragraph, text, base_size=11, base_color=DARK_GRAY):
    """Process bold (**text**) and other inline markdown formatting."""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            set_run_font(run, size=base_size, bold=True, color=base_color)
        else:
            if part:
                run = paragraph.add_run(part)
                set_run_font(run, size=base_size, bold=False, color=base_color)


def main():
    print("Reading markdown file...")
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    print(f"  Read {len(lines)} lines")

    # First pass: extract day headings for TOC
    days = []
    seen_days = set()
    for line in lines:
        line_stripped = line.strip()
        # Pattern 1: # DIA 1 — QUARTA-FEIRA 01/04/2026
        m = re.match(r'^#\s+DIA\s+(\d+)\s*[—–-]\s*(.*)', line_stripped)
        if not m:
            # Pattern 2: # ║           DIA 6 — SEGUNDA-FEIRA 06/04/2026              ║
            m = re.match(r'^#\s*║\s*DIA\s+(\d+)\s*[—–-]\s*(.*?)\s*║?\s*$', line_stripped)
        if m:
            day_num = m.group(1)
            if day_num not in seen_days:
                seen_days.add(day_num)
                day_text = m.group(2).strip().rstrip('║').strip()
                days.append(f"Dia {day_num} — {day_text}")

    print(f"  Found {len(days)} day entries for TOC")

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Add cover page and TOC
    add_cover_page(doc)
    add_toc_page(doc, days)

    # Second pass: process content
    print("Processing content...")
    in_code_block = False
    skip_initial_header = True  # skip the very first lines (we made a cover page)
    lines_after_header = 0

    i = 0
    total = len(lines)
    progress_step = total // 10

    while i < total:
        if progress_step > 0 and i % progress_step == 0:
            pct = int(i / total * 100)
            print(f"  {pct}% processed...")

        line = lines[i].rstrip('\n').rstrip('\r')
        stripped = line.strip()
        i += 1

        # Skip the initial file header (already in cover page)
        if skip_initial_header:
            if stripped.startswith('# ====') or (stripped.startswith('# DIA ') and '—' in stripped) or ('║' in stripped and 'DIA' in stripped):
                skip_initial_header = False
            else:
                continue

        # Handle code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # Starting code block - add a subtle indicator
                pass
            continue

        if in_code_block:
            # Render code block content with monospace-like smaller font
            if stripped:
                # Check if it's a section divider line
                if re.match(r'^[=]+$', stripped):
                    add_styled_paragraph(doc, "—" * 40, size=9, color=RGBColor(0x99, 0x99, 0x99), space_after=2, space_before=2)
                elif stripped.startswith('HOOK') or stripped.startswith('CTA') or stripped.startswith('CONTEXTO') or stripped.startswith('EXEMPLO') or stripped.startswith('CONTEUDO') or stripped.startswith('ENCERRAMENTO'):
                    # Section headers inside scripts - gold bold
                    # Clean up the text (remove parenthetical timing info for cleaner display)
                    add_styled_paragraph(doc, stripped, size=11, bold=True, color=GOLD, space_before=8, space_after=4)
                elif stripped.startswith('ALE FALA:') or stripped.startswith('TEXTO NA TELA'):
                    add_styled_paragraph(doc, stripped, size=11, bold=True, color=NAVY, space_after=2)
                elif stripped.startswith('INSTRUCAO'):
                    add_styled_paragraph(doc, stripped, size=10, bold=True, color=RGBColor(0x66, 0x66, 0x66), space_after=2)
                elif stripped.startswith('PILULA VISUAL'):
                    add_styled_paragraph(doc, stripped, size=10, bold=False, color=RGBColor(0x66, 0x66, 0x66), space_after=1)
                elif stripped.startswith('#'):
                    # Hashtags inside captions
                    add_styled_paragraph(doc, stripped, size=10, bold=False, color=RGBColor(0x55, 0x77, 0xAA), space_after=1)
                elif stripped.startswith('-') or stripped.startswith('•'):
                    add_styled_paragraph(doc, stripped, size=10, color=DARK_GRAY, space_after=1)
                else:
                    add_styled_paragraph(doc, stripped, size=10, color=DARK_GRAY, space_after=2)
            else:
                add_styled_paragraph(doc, "", size=6, space_after=2)
            continue

        # --- (horizontal rule)
        if re.match(r'^-{3,}$', stripped):
            # Add a light separator
            add_styled_paragraph(doc, "", size=4, space_after=2)
            continue

        # Empty line
        if not stripped:
            continue

        # Decorative lines (════ or ===== or ╔╗╚╝║ box drawing at heading level)
        if re.match(r'^#\s*[═=╔╗╚╝╠╣─]{5,}', stripped):
            continue
        # Box drawing lines without # prefix
        if re.match(r'^[═╔╗╚╝╠╣─║]{3,}$', stripped):
            continue

        # H1 headings
        if stripped.startswith('# '):
            text = stripped[2:].strip()

            # Box-drawing DIA headers: ║           DIA 6 — SEGUNDA-FEIRA 06/04/2026              ║
            m_box = re.match(r'║\s*DIA\s+(\d+)\s*[—–-]\s*(.*?)\s*║?\s*$', text)
            if m_box:
                clean_text = f"DIA {m_box.group(1)} — {m_box.group(2).strip().rstrip('║').strip()}"
                doc.add_page_break()
                add_styled_paragraph(doc, clean_text, size=18, bold=True, color=NAVY, space_after=6, space_before=12)
                continue

            # SEMANA headers
            if 'SEMANA' in text and ('—' in text or '-' in text):
                doc.add_page_break()
                add_styled_paragraph(doc, text, size=20, bold=True, color=NAVY,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, space_before=24)
                continue

            # DIA headers (standard format)
            m = re.match(r'DIA\s+\d+', text)
            if m:
                doc.add_page_break()
                add_styled_paragraph(doc, text, size=18, bold=True, color=NAVY, space_after=6, space_before=12)
                continue

            # POST headers (REEL or CARROSSEL)
            if text.startswith('POST') or text.startswith('"'):
                is_reel = 'REEL' in text.upper()
                is_carousel = 'CARROSSEL' in text.upper() or 'CARROS' in text.upper()

                if text.startswith('POST'):
                    color = GOLD if is_reel else NAVY
                    add_styled_paragraph(doc, "", size=6, space_after=4)
                    add_styled_paragraph(doc, text, size=14, bold=True, color=color, space_after=2, space_before=16)
                else:
                    # Title line (in quotes)
                    add_styled_paragraph(doc, text, size=13, bold=True, color=DARK_GRAY, space_after=8)
                continue

            # ROTEIROS COMPLETOS header
            if 'ROTEIROS COMPLETOS' in text:
                add_styled_paragraph(doc, text, size=14, bold=True, color=NAVY, space_after=6, space_before=6)
                continue

            # Generic H1
            add_styled_paragraph(doc, text, size=14, bold=True, color=NAVY, space_after=6, space_before=8)
            continue

        # H2 headings
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            add_styled_paragraph(doc, text, size=13, bold=True, color=NAVY, space_after=6, space_before=6)
            continue

        # H3 headings
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            # Special styling for key sections
            if 'LEGENDA' in text.upper():
                add_styled_paragraph(doc, text, size=12, bold=True, color=GOLD, space_after=6, space_before=10)
            elif 'ROTEIRO' in text.upper():
                add_styled_paragraph(doc, text, size=12, bold=True, color=GOLD, space_after=6, space_before=10)
            elif 'SLIDES' in text.upper():
                add_styled_paragraph(doc, text, size=12, bold=True, color=GOLD, space_after=6, space_before=10)
            else:
                add_styled_paragraph(doc, text, size=12, bold=True, color=NAVY, space_after=6, space_before=8)
            continue

        # H4+ headings
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            add_styled_paragraph(doc, text, size=11, bold=True, color=NAVY, space_after=4, space_before=6)
            continue

        # Table rows (markdown tables)
        if stripped.startswith('|'):
            # Skip separator rows like |---|---|
            if re.match(r'^\|[\s\-|]+\|$', stripped):
                continue
            # Render table rows as tab-separated text
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            row_text = "   |   ".join(cells)
            add_styled_paragraph(doc, row_text, size=9, color=DARK_GRAY, space_after=1)
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Cm(0.5)
            # Add bullet character
            run = p.add_run("  •  ")
            set_run_font(run, size=10, color=GOLD, bold=True)
            # Process inline formatting for the rest
            process_inline_formatting(p, text, base_size=10, base_color=DARK_GRAY)
            continue

        # Bold-start lines (like **Tipo de conteudo:** ...)
        if stripped.startswith('**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(1)
            process_inline_formatting(p, stripped, base_size=10, base_color=DARK_GRAY)
            continue

        # Regular text
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        process_inline_formatting(p, stripped, base_size=10, base_color=DARK_GRAY)

    # Save
    print("Saving document...")
    doc.save(OUTPUT_FILE)
    file_size = os.path.getsize(OUTPUT_FILE)
    print(f"Done! File saved to: {OUTPUT_FILE}")
    print(f"File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")


if __name__ == '__main__':
    main()
