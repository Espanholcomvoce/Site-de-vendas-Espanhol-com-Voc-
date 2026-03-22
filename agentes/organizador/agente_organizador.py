"""
Agente Organizador — @espanholcomvoce
======================================
Le os roteiros JSON e gera um .docx por roteiro, organizado por mes,
salvo no Google Drive em G:\Meu Drive\EspanholComVoce\Conteudo2026\.

Estrutura de saida:
  G:\Meu Drive\EspanholComVoce\Conteudo2026\
    03_Marco\
      2026-03-20_REEL_Viajante_se-vira-nos-30.docx
      2026-03-20_CARROSSEL_VaiMorarFora_metodo-bastidor.docx
    04_Abril\
      ...

Cada .docx tem:
  - Titulo com data, horario e tipo
  - Secoes: Hook, Pilulas, Fala da Ale, CTA, Legenda
  - Formatado para leitura como teleprompter (fonte grande, espaçado)
"""

import json
import re
import unicodedata
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Caminhos ---
BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "output"
DRIVE_BASE = Path("G:/Meu Drive/EspanholComVoce/Conteudo2026")

MESES_PT = {
    1: "01_Janeiro", 2: "02_Fevereiro", 3: "03_Marco",
    4: "04_Abril", 5: "05_Maio", 6: "06_Junho",
    7: "07_Julho", 8: "08_Agosto", 9: "09_Setembro",
    10: "10_Outubro", 11: "11_Novembro", 12: "12_Dezembro",
}


def slugify(texto, max_len=50):
    """Converte texto em slug seguro para nome de arquivo."""
    # Remove acentos
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    # Minusculas, substitui espacos e caracteres especiais
    texto = texto.lower().strip()
    texto = re.sub(r"[^\w\s-]", "", texto)
    texto = re.sub(r"[\s_]+", "-", texto)
    texto = re.sub(r"-+", "-", texto).strip("-")
    return texto[:max_len]


# ================================================================
# ESTILOS DO DOCUMENTO
# ================================================================

def configurar_estilos(doc):
    """Configura estilos base do documento para leitura como teleprompter."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(14)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5


def add_titulo_principal(doc, texto):
    """Titulo grande no topo do documento."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 30, 56)  # Azul Marinho


def add_metadados(doc, roteiro):
    """Bloco de metadados do post."""
    dados = [
        f"Data: {roteiro['data']} ({roteiro['dia_semana']}) — {roteiro['horario']}",
        f"Formato: {roteiro['formato']} | Funil: {roteiro['funil']}",
        f"Tipo: {roteiro['tipo_conteudo']} (ID {roteiro['tipo_conteudo_id']})",
        f"Avatar: {roteiro['avatar']}",
        f"Keyword ManyChat: {roteiro['keyword_manychat']}",
    ]
    for linha in dados:
        p = doc.add_paragraph()
        run = p.add_run(linha)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()  # Espaco


def add_secao(doc, titulo, cor_hex=None):
    """Adiciona titulo de secao."""
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(16)
    if cor_hex:
        r, g, b = int(cor_hex[1:3], 16), int(cor_hex[3:5], 16), int(cor_hex[5:7], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    else:
        run.font.color.rgb = RGBColor(0, 30, 56)


def add_fala_teleprompter(doc, texto, destaque=False):
    """Adiciona texto de fala em formato teleprompter (grande, legivel)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(texto)
    run.font.size = Pt(18) if destaque else Pt(16)
    run.bold = destaque
    if destaque:
        run.font.color.rgb = RGBColor(255, 66, 28)  # Vermelho Laranja


def add_texto_normal(doc, texto):
    """Texto normal."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(13)


def add_instrucao(doc, texto):
    """Instrucao de gravacao (italico, cinza)."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{texto}]")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)


def add_separador(doc):
    """Linha separadora entre secoes."""
    p = doc.add_paragraph()
    run = p.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(200, 200, 200)


def add_pilula_visual(doc, num, texto, fundo, texto_cor):
    """Representa uma pilula visual no documento."""
    p = doc.add_paragraph()
    label = p.add_run(f"  Pilula {num}: ")
    label.font.size = Pt(12)
    label.bold = True
    label.font.color.rgb = RGBColor(100, 100, 100)
    conteudo = p.add_run(f" {texto} ")
    conteudo.bold = True
    conteudo.font.size = Pt(14)
    if "branco" in texto_cor.lower():
        conteudo.font.color.rgb = RGBColor(0, 30, 56)
    else:
        conteudo.font.color.rgb = RGBColor(0, 0, 0)
    fundo_info = p.add_run(f"  (fundo: {fundo})")
    fundo_info.font.size = Pt(10)
    fundo_info.font.color.rgb = RGBColor(150, 150, 150)


# ================================================================
# GERAR DOCX — REEL
# ================================================================

def gerar_docx_reel(roteiro, caminho):
    """Gera .docx de um reel formatado como teleprompter."""
    doc = Document()
    configurar_estilos(doc)

    # Titulo
    tema_curto = roteiro["tema"][:80]
    add_titulo_principal(doc, f"REEL — {tema_curto}")
    add_metadados(doc, roteiro)

    rot = roteiro["roteiro"]

    # --- HOOK ---
    add_separador(doc)
    add_secao(doc, "HOOK (3 primeiros segundos)", "#ff421c")
    add_fala_teleprompter(doc, rot["hook_3s"]["fala"], destaque=True)
    add_instrucao(doc, rot["hook_3s"]["instrucao_gravacao"])

    # --- PILULAS VISUAIS ---
    add_separador(doc)
    add_secao(doc, "PILULAS VISUAIS (overlay no video)")
    pilulas = rot["pilulas_visuais"]
    for num, key in enumerate(["pilula_1", "pilula_2", "pilula_3"], 1):
        p = pilulas[key]
        add_pilula_visual(doc, num, p["texto"], p["fundo"], p["texto_cor"])

    # --- FALA DA ALE ---
    add_separador(doc)
    add_secao(doc, "FALA DA ALE (teleprompter)", "#001e38")
    dev = rot["desenvolvimento"]
    for i, bloco in enumerate(dev["blocos_fala"], 1):
        add_fala_teleprompter(doc, f"{bloco}")
        # Linha vazia entre blocos para respiracao
        doc.add_paragraph()

    # Instrucoes de gravacao
    add_secao(doc, "Instrucoes de Gravacao", "#79458f")
    for instrucao in dev["instrucoes_gravacao"]:
        add_instrucao(doc, instrucao)

    # Palavras em espanhol
    if dev.get("palavras_em_espanhol") and "verificar" not in dev["palavras_em_espanhol"].lower():
        add_texto_normal(doc, f"Palavras-chave em espanhol: {dev['palavras_em_espanhol']}")

    # --- CTA FINAL ---
    add_separador(doc)
    add_secao(doc, "CTA FINAL", "#ff421c")
    add_fala_teleprompter(doc, rot["cta_final"]["fala"], destaque=True)
    add_instrucao(doc, rot["cta_final"]["instrucao_gravacao"])

    # --- LEGENDA ---
    add_separador(doc)
    add_secao(doc, "LEGENDA (copiar e colar)")
    # Legenda em fonte menor, pronta pra copiar
    p = doc.add_paragraph()
    run = p.add_run(roteiro["legenda"])
    run.font.size = Pt(11)
    run.font.name = "Arial"

    doc.save(str(caminho))


# ================================================================
# GERAR DOCX — CARROSSEL
# ================================================================

def gerar_docx_carrossel(roteiro, caminho):
    """Gera .docx de um carrossel com slides detalhados."""
    doc = Document()
    configurar_estilos(doc)

    # Titulo
    tema_curto = roteiro["tema"][:80]
    add_titulo_principal(doc, f"CARROSSEL — {tema_curto}")
    add_metadados(doc, roteiro)

    rot = roteiro["roteiro"]

    # --- IDENTIDADE VISUAL ---
    add_separador(doc)
    add_secao(doc, "IDENTIDADE VISUAL")
    iv = rot["identidade_visual"]
    add_texto_normal(doc, f"Cor de fundo: {iv['cor_nome']} ({iv['cor_fundo']})")
    add_texto_normal(doc, f"Fonte titulo: {iv['fonte_titulo']}")
    add_texto_normal(doc, f"Logo: {iv['logo']}")
    add_texto_normal(doc, f"Icone: {iv['icone_mao']}")

    # --- SLIDES ---
    add_separador(doc)
    add_secao(doc, f"SLIDES ({roteiro['total_slides']} slides)", "#001e38")

    for slide in rot["slides"]:
        doc.add_paragraph()
        # Numero e tipo do slide
        p = doc.add_paragraph()
        header = p.add_run(f"SLIDE {slide['slide']} — {slide['tipo_slide']}")
        header.bold = True
        header.font.size = Pt(14)
        header.font.color.rgb = RGBColor(0, 30, 56)

        # Cor de fundo do slide
        cor = slide.get("cor_fundo", "#FFFFFF")
        p_cor = doc.add_paragraph()
        run_cor = p_cor.add_run(f"  Fundo: {cor}")
        run_cor.font.size = Pt(10)
        run_cor.font.color.rgb = RGBColor(150, 150, 150)

        # Texto principal (grande, como teleprompter)
        texto_principal = slide["texto_principal"].replace("\\n", "\n")
        for linha in texto_principal.split("\n"):
            if linha.strip():
                add_fala_teleprompter(doc, linha.strip(), destaque=(slide["tipo_slide"] == "CTA"))

        # Texto secundario
        texto_sec = slide.get("texto_secundario", "")
        if texto_sec:
            texto_sec = texto_sec.replace("\\n", "\n")
            for linha in texto_sec.split("\n"):
                if linha.strip():
                    add_texto_normal(doc, linha.strip())

        add_separador(doc)

    # --- LEGENDA ---
    add_secao(doc, "LEGENDA (copiar e colar)")
    p = doc.add_paragraph()
    run = p.add_run(roteiro["legenda"])
    run.font.size = Pt(11)
    run.font.name = "Arial"

    doc.save(str(caminho))


# ================================================================
# CLASSE PRINCIPAL
# ================================================================

class AgenteOrganizador:
    """Organiza roteiros em .docx no Google Drive."""

    def __init__(self, arquivo_roteiros, drive_base=None):
        self.arquivo_roteiros = Path(arquivo_roteiros)
        self.drive_base = Path(drive_base) if drive_base else DRIVE_BASE

    def _nome_arquivo(self, roteiro):
        """Gera nome do arquivo: DATA_FORMATO_Avatar_tema-slug.docx"""
        data = roteiro["data"]
        horario = roteiro["horario"].replace(":", "")
        formato = roteiro["formato"]
        avatar = roteiro["avatar"].replace(" ", "")
        tema_slug = slugify(roteiro["tema"])
        return f"{data}_{horario}_{formato}_{avatar}_{tema_slug}.docx"

    def _pasta_mes(self, data_str):
        """Retorna pasta do mes: 03_Marco, 04_Abril, etc."""
        data = datetime.strptime(data_str, "%Y-%m-%d")
        return MESES_PT.get(data.month, f"{data.month:02d}")

    def executar(self):
        print(f"\n  Agente Organizador — @espanholcomvoce")
        print(f"  Lendo roteiros: {self.arquivo_roteiros.name}")
        print(f"  Destino: {self.drive_base}\n")

        with open(self.arquivo_roteiros, "r", encoding="utf-8") as f:
            roteiros = json.load(f)

        total = len(roteiros)
        reels = 0
        carrosseis = 0
        erros = 0
        meses_criados = set()

        for i, roteiro in enumerate(roteiros, 1):
            try:
                # Determina pasta do mes
                pasta_mes = self._pasta_mes(roteiro["data"])
                meses_criados.add(pasta_mes)
                pasta = self.drive_base / pasta_mes
                pasta.mkdir(parents=True, exist_ok=True)

                # Nome do arquivo
                nome = self._nome_arquivo(roteiro)
                caminho = pasta / nome

                # Gera docx
                if roteiro["formato"] == "REEL":
                    gerar_docx_reel(roteiro, caminho)
                    reels += 1
                else:
                    gerar_docx_carrossel(roteiro, caminho)
                    carrosseis += 1

                # Progresso a cada 50
                if i % 50 == 0 or i == total:
                    print(f"  [{i}/{total}] processados...")

            except Exception as e:
                erros += 1
                print(f"  [ERRO] Post {roteiro.get('post_id', '?')}: {e}")

        # Resumo
        print(f"\n  {'=' * 55}")
        print(f"  ORGANIZADOR CONCLUIDO")
        print(f"  {'=' * 55}")
        print(f"  Total: {total} roteiros processados")
        print(f"  Reels: {reels} .docx")
        print(f"  Carrosseis: {carrosseis} .docx")
        print(f"  Erros: {erros}")
        print(f"  Meses: {', '.join(sorted(meses_criados))}")
        print(f"  Destino: {self.drive_base}")
        print(f"  {'=' * 55}")


# --- CLI ---
# Uso: python agente_organizador.py <roteiros.json> [caminho_drive]
# Padrao: usa o roteiro mais recente e G:\Meu Drive\EspanholComVoce\Conteudo2026
if __name__ == "__main__":
    import sys
    import glob

    if len(sys.argv) >= 2:
        arquivo = sys.argv[1]
    else:
        padrao = str(OUTPUT_DIR / "roteiros_2*.json")
        # Exclui os resumos
        arquivos = [f for f in sorted(glob.glob(padrao)) if "resumo" not in f]
        if not arquivos:
            print("  [ERRO] Nenhum arquivo de roteiros encontrado em agentes/output/")
            print("  Execute primeiro o Agente Roteirista.")
            sys.exit(1)
        arquivo = arquivos[-1]
        print(f"  Usando roteiros mais recentes: {arquivo}")

    drive = sys.argv[2] if len(sys.argv) >= 3 else None

    agente = AgenteOrganizador(arquivo, drive)
    agente.executar()
