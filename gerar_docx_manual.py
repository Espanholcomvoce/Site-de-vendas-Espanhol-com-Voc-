"""
Compila todos os arquivos do Manual do Método em um único .docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, json

doc = Document()

# Estilos
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level, size in [('Heading 1', 22), ('Heading 2', 16), ('Heading 3', 13)]:
    s = doc.styles[level]
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
    s.font.bold = True


def divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━' * 60)
    r.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)
    r.font.size = Pt(8)


def section_label(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)


def add_md_content(filepath):
    """Read a markdown file and add its content with basic formatting."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        stripped = line.rstrip('\n')

        if stripped.startswith('# ') and not stripped.startswith('##'):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('---'):
            divider()
        elif stripped.strip() == '':
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            # Handle bold markers
            text = stripped
            if text.startswith('**') and '**' in text[2:]:
                # Bold line
                parts = text.split('**')
                for i, part in enumerate(parts):
                    if part.strip():
                        r = p.add_run(part)
                        r.font.size = Pt(11)
                        if i % 2 == 1:
                            r.font.bold = True
            else:
                r = p.add_run(text)
                r.font.size = Pt(11)


def add_json_summary(filepath, title):
    """Read a JSON config and add a formatted summary."""
    with open(filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc.add_heading(title, level=2)
    # Pretty print the JSON
    formatted = json.dumps(data, indent=2, ensure_ascii=False)
    p = doc.add_paragraph()
    r = p.add_run(formatted)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)


# ═══════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ESPANHOL COM VOCÊ')
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('@espanholcomvoce')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0xfb, 0xbf, 0x24)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MANUAL DO MÉTODO')
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Pacote Completo de Implementação')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Texto do PDF • Emails • Scripts de Vídeo • Fluxo ManyChat • Configs dos Agentes')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Março 2026')
r.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════
doc.add_heading('ÍNDICE', level=1)
doc.add_paragraph()

indice_items = [
    ('PARTE 1 — TEXTO COMPLETO DO PDF', True),
    ('  Capa e abertura', False),
    ('  Capítulo 1: Erro #1 — Padrões, não palavras', False),
    ('  Capítulo 2: Erro #2 — A ordem invertida', False),
    ('  Capítulo 3: Erro #3 — Calibração', False),
    ('  Capítulo 4: Seus próximos 3 passos', False),
    ('  Capítulo 5: O Método Imersão Nativa', False),
    ('', False),
    ('PARTE 2 — SEQUÊNCIA DE 5 EMAILS', True),
    ('  Email 1: Entrega do Manual (dia 0)', False),
    ('  Email 2: Reforço Erro #1 (dia 2)', False),
    ('  Email 3: Prova social (dia 4)', False),
    ('  Email 4: Transição para VSL (dia 6)', False),
    ('  Email 5: Oferta (dia 8)', False),
    ('', False),
    ('PARTE 3 — SCRIPTS DOS 4 MINI-VÍDEOS', True),
    ('  Vídeo 1: Minha história (60s)', False),
    ('  Vídeo 2: 3 expressões nativas (45s)', False),
    ('  Vídeo 3: Entender vs Falar (45s)', False),
    ('  Vídeo 4: Progressão calibrada (60s)', False),
    ('', False),
    ('PARTE 4 — FLUXO MANYCHAT', True),
    ('  Trigger, captura, tags, automação', False),
    ('', False),
    ('PARTE 5 — CONFIGURAÇÕES DOS AGENTES', True),
    ('  horarios.json (atualizado)', False),
    ('  objetivos.json (atualizado)', False),
    ('  tipos_conteudo.json (atualizado)', False),
]

for text, is_bold in indice_items:
    if not text:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.bold = is_bold
        p.runs[0].font.size = Pt(11 if is_bold else 10)
        if not is_bold:
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════
# PARTE 1 — TEXTO DO PDF
# ═══════════════════════════════════════════
doc.add_heading('PARTE 1 — TEXTO COMPLETO DO PDF', level=1)
p = doc.add_paragraph('22 páginas — pronto para design. Inclui marcadores [PÁGINA X], [DIAGRAMA], [QR CODE] e [FOTO] para o designer.')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
divider()
doc.add_paragraph()

add_md_content('D:/EspanholComVoce/iscas/manual_do_metodo/texto_completo.md')

doc.add_page_break()

# ═══════════════════════════════════════════
# PARTE 2 — EMAILS
# ═══════════════════════════════════════════
doc.add_heading('PARTE 2 — SEQUÊNCIA DE 5 EMAILS', level=1)
p = doc.add_paragraph('Disparar via Mailchimp. Tag: lead_manual_metodo. Automação: dias 0, 2, 4, 6, 8.')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
divider()
doc.add_paragraph()

add_md_content('D:/EspanholComVoce/iscas/manual_do_metodo/emails_sequencia.md')

doc.add_page_break()

# ═══════════════════════════════════════════
# PARTE 3 — SCRIPTS DE VÍDEO
# ═══════════════════════════════════════════
doc.add_heading('PARTE 3 — SCRIPTS DOS 4 MINI-VÍDEOS', level=1)
p = doc.add_paragraph('Ale grava 4 vídeos curtos (total ~3,5 minutos). Linkados via QR code dentro do PDF.')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
divider()
doc.add_paragraph()

add_md_content('D:/EspanholComVoce/iscas/manual_do_metodo/scripts_videos.md')

doc.add_page_break()

# ═══════════════════════════════════════════
# PARTE 4 — MANYCHAT
# ═══════════════════════════════════════════
doc.add_heading('PARTE 4 — FLUXO MANYCHAT', level=1)
p = doc.add_paragraph('Especificação completa do fluxo de automação para keyword MANUAL.')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
divider()
doc.add_paragraph()

add_md_content('D:/EspanholComVoce/iscas/manual_do_metodo/manychat_flow.md')

doc.add_page_break()

# ═══════════════════════════════════════════
# PARTE 5 — CONFIGS
# ═══════════════════════════════════════════
doc.add_heading('PARTE 5 — CONFIGURAÇÕES DOS AGENTES', level=1)
p = doc.add_paragraph('Arquivos JSON atualizados para o sistema de agentes de IA — Abril 2026.')
p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
divider()

add_json_summary('D:/EspanholComVoce/agentes/config/horarios.json', 'horarios.json')
doc.add_page_break()
add_json_summary('D:/EspanholComVoce/agentes/config/objetivos.json', 'objetivos.json')
doc.add_page_break()
add_json_summary('D:/EspanholComVoce/agentes/config/tipos_conteudo.json', 'tipos_conteudo.json')

# ═══════════════════════════════════════════
# RODAPÉ
# ═══════════════════════════════════════════
doc.add_page_break()
divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nDocumento gerado em Março 2026\n')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Espanhol com Você — @espanholcomvoce\n')
r.font.size = Pt(12)
r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Manual do Método — Pacote Completo de Implementação')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# Salvar
output = r'D:\EspanholComVoce\manual_do_metodo_completo.docx'
doc.save(output)
print(f'Salvo: {output}')
print(f'Tamanho: {os.path.getsize(output) / 1024:.0f} KB')
